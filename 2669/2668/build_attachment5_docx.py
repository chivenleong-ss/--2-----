from pathlib import Path

from docx import Document


SRC = Path("附件5：数字化审计模型-数字化营销审计智能分析与穿透预警模型6_完整版修订.md")
OUT = Path("附件5：数字化审计模型-数字化营销审计智能分析与穿透预警模型6_完整版修订.docx")


def add_paragraph(doc: Document, text: str) -> None:
    stripped = text.strip()
    if not stripped:
        doc.add_paragraph("")
        return
    if stripped.startswith("# "):
        doc.add_heading(stripped[2:].strip(), level=0)
        return
    if stripped.startswith("## "):
        doc.add_heading(stripped[3:].strip(), level=1)
        return
    if stripped.startswith("### "):
        doc.add_heading(stripped[4:].strip(), level=2)
        return
    doc.add_paragraph(stripped)


def main() -> None:
    content = SRC.read_text(encoding="utf-8").splitlines()
    doc = Document()
    for line in content:
        add_paragraph(doc, line)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
