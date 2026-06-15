"""
Final Word document update: DMP-only reality.
Remove all 附表6/客户台账 references. Add DMP-only field classification.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import time

SRC = r'C:\Users\sasa\Desktop\模型建设\模型2：市场营销\【完善版修订】全面数字化营销审计系统实施方案_v2.9_完整版_173519.docx'

doc = Document(SRC)
body = doc.element.body

def insert_para_before(ref, text, heading_level=None, bold=False):
    p = OxmlElement('w:p')
    if heading_level:
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), f'Heading{heading_level}')
        pPr.append(pStyle)
        p.append(pPr)
    r = OxmlElement('w:r')
    if bold or heading_level:
        rPr = OxmlElement('w:rPr')
        b = OxmlElement('w:b')
        rPr.append(b)
    r.append(rPr) if (bold or heading_level) else None
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    ref.addprevious(p)

# === 1. Title & P6 ===
doc.paragraphs[0].text = '全面数字化营销审计系统端到端落地实施方案（完善版修订稿 v2.9 - DMP数据驱动版）'
doc.paragraphs[1].text = '**修订日期**：2026-06-02'
doc.paragraphs[6].text = (
    'v2.9核心更新：（1）DMP四节点时间线校验替代OA审批流；'
    '（2）模型2.1新增付款条件不达标制度交叉校验；'
    '（3）签约报量(四局)170字段+DMP中标报量152字段为唯一数据源；'
    '（4）维度三全面重构——附表6/客户台账不可获取，以DMP中标报量+制度内嵌名单(手册0520附件5/10)替代；'
    '（5）offline_fields调整为51字段(移除客户台账6字段)；'
    '（6）SAP 4日期字段为唯一不可替代的🔴高优缺口。'
)

# === 2. P11 data source ===
doc.paragraphs[11].text = (
    '说明：DMP线上导出数据为唯一数据源——签约报量(四局)170字段为主表，'
    '中标报量152字段补充投标阶段时间线及金额字段。'
    '审计附表(手工填列)仅补充DMP中没有的13个财务实际数据字段。'
    '区域认定按年度分表。'
    '附表6(战略客户统计表)和客户台账(拜访/评级数据)不可获取——'
    '模型1.3/3.1/3.2/3.3已改用DMP中标报量+制度内嵌名单替代。'
)

# === 3. P13 ===
doc.paragraphs[13].text = (
    '从DMP系统170字段(签约报量四局)+中标报量152字段中，'
    '按审计模型实际需要用到的字段梳理如下。'
    '维度三(客户健康度)不再依赖附表6/客户台账——全部客户侧检测基于DMP中标报量+签约报量完成。'
)

# === 4. Section 2.3 heading ===
for i, p in enumerate(doc.paragraphs):
    if p.style.name.startswith('Heading 3') and '2.3 客户属性类' in p.text:
        p.text = '2.3 客户属性类（模型1.3 / 3.1 / 3.2 / 3.3）[1.3已内嵌制度名单2026]'
        break

# === 5. Insert sections 2.11-2.13 before section 3 ===
target_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.style.name.startswith('Heading 2') and '3. 动态阈值配置文件' in p.text:
        target_idx = i
        break

if target_idx:
    ref = doc.paragraphs[target_idx]._element

    # 2.11
    insert_para_before(ref, '', None)
    insert_para_before(ref, '2.11 审计附表补充字段（仅附表1可用）', heading_level=3)
    insert_para_before(ref, '以下为审计附表1中被各模型消费的13个财务实际数据字段。附表3/附表6不可获取——模型3.1/1.3/3.2已用DMP+制度名单替代。', None)
    insert_para_before(ref, '', None)
    items_app1 = [
        '【全部模型】项目编码（财务部）— JOIN键',
        '【前置+2.2】项目状态 — 停工/退场/停缓建识别',
        '【前置+2.2】实际完成产值 — 180天产值检测 + 在施验证',
        '【前置+2.2+2.3】累计收款 — 360天收款检测 + 保证金回收佐证',
        '【2.1】应收未收款 — 履约监控分级检测',
        '【2.2】最近一期成本分析利润率 — A值偏差计算',
        '【2.2】未开工或退场原因 — 停工退场根因补充',
        '【2.3】预收款应收款+实收款 — 预收款缺口计算',
        '【2.3】资金结余+负流原因分析 — 负现金流检测',
        '【2.4】三证取得情况明细 — DMP仅有Y/N标记，附表补充实际状态',
        '【2.4】工期延误天数 — 实际履约偏差定量',
    ]
    for item in items_app1:
        insert_para_before(ref, item, None)
    insert_para_before(ref, '', None)
    insert_para_before(ref,
        '附表3(中标率表)和附表6(战略客户统计表)不可获取。'
        '模型3.1以DMP中标→签约转化率替代附表3中标率；'
        '模型1.3/3.2以制度内嵌名单(手册0520附件5/10,2026年)替代附表6客户分级和主责维护。', None)
    insert_para_before(ref, '', None)

    # 2.12
    insert_para_before(ref, '2.12 中标报量补充字段（8个，已LEFT JOIN入统一DataFrame）', heading_level=3)
    insert_para_before(ref, '', None)
    items_bid = [
        '【1.4】中标额(元) — 中标签约金额交叉验证',
        '【1.4】中标报量时间 + 预计签约时间 — 时间线校验补充节点',
        '【1.2+3.3】是否高端客户/市场/项目 — 旧版三高标记，与签约报量优质标记交叉校验',
        '【1.2】项目分类 + 项目模式类 — 业务结构分析补充维度',
        '【3.1+3.3】中标额+中标报量时间 → 中标转化率 + 中标未签约僵尸检测',
    ]
    for item in items_bid:
        insert_para_before(ref, item, None)
    insert_para_before(ref, '', None)

    # 2.13
    insert_para_before(ref, '2.13 SAP/财务系统待采集字段（7个，唯一外部依赖）', heading_level=3)
    insert_para_before(ref, '以下为当前唯一仍需从外部系统获取的字段（采集方：财务部）。客户台账已通过DMP中标报量+制度名单覆盖。', None)
    insert_para_before(ref, '', None)
    items_sap = [
        '【2.3】保证金约定退还日期 — 逾期天数计算 (>90天重大风险)',
        '【2.3】保证金实际回收日期 — 对比约定退还日计算逾期',
        '【2.3】预收款约定支付日期 — 预收款逾期天数计算',
        '【2.3】预收款实际收款日期 — 对比约定支付日计算逾期',
        '【2.2】实际产值(SAP口径) — 转化率计算',
        '【2.1】实际收款总额(SAP口径) — 签约后收款进度监控',
        '【2.1】已结算未收款(SAP口径) — 履约监控分级检测',
    ]
    for item in items_sap:
        insert_para_before(ref, item, None)
    insert_para_before(ref, '', None)
    insert_para_before(ref,
        'SAP 4个日期字段为模型2.3核心依赖——缺失导致保证金/预收款逾期检测完全失效。'
        '这是当前唯一不可替代的外部数据缺口(共4个高优字段)。', None)

print('Phase 1-5: Title, data source, field classification done')

# === 6. Chapter 6 updates ===
for i, p in enumerate(doc.paragraphs):
    if p.style.name.startswith('Heading 3') and '6.2.2 中标率表' in p.text:
        p.text = '6.2.2 中标率表(附表3) —— 审计组 (v2.9:不可获取，模型3.1已用DMP中标转化率替代)'
    if p.style.name.startswith('Heading 3') and '6.2.3 战略客户表' in p.text:
        p.text = '6.2.3 战略客户表(附表6) —— 审计组 (v2.9:不可获取，模型1.3/3.2已用制度内嵌名单替代)'
    if p.style.name.startswith('Heading 3') and '6.2.5 市场部客户台账' in p.text:
        p.text = '6.2.5 市场部客户台账 —— 市场部 (v2.9:不可获取，DMP中标报量+制度名单已覆盖维度三)'

# P868
for i, p in enumerate(doc.paragraphs):
    if '客户台账共需' in p.text and '6个字段' in p.text:
        p.text = (
            '客户台账共需6个字段，当前不可获取。v2.9维度三已全面重构：'
            '模型3.1以DMP中标报量-签约报量转化率替代拜访维度；'
            '模型3.3以中标-签约交叉校验+优质标记校验替代评级核查；'
            '模型3.2以制度内嵌名单(手册0520附件5)替代战略客户识别。'
            '维度三3个模型均可在无台账情况下全功能运行。'
        )

# P878
for i, p in enumerate(doc.paragraphs):
    if '当前最急缺' in p.text and 'SAP' in p.text and '客户台账' in p.text:
        p.text = (
            '**当前最急缺(v2.9)**：SAP 4个日期字段(模型2.3核心)——'
            '保证金/预收款逾期检测完全依赖此数据源，DMP无替代字段。'
            '这是当前唯一不可替代的高优外部数据缺口。'
            'OA审批流DMP四节点已覆盖80%场景；客户台账DMP中标报量已覆盖全部维度三需求。'
            '共 **4个高优字段**(SAP日期x4)。'
        )
        break

print('Phase 6: Chapter 6 done')

# === 7. Appendix B ===
for i, p in enumerate(doc.paragraphs):
    if 'B. 市场部客户台账' in p.text and p.style.name.startswith('Heading'):
        if i + 1 < len(doc.paragraphs):
            doc.paragraphs[i + 1].text = (
                'v2.9：客户台账不可获取。维度三(客户健康度)三个模型已通过DMP中标报量+'
                '制度内嵌名单(手册0520附件5/10,2026年)实现全功能运行：\n'
                '- 模型3.1：中标→签约转化率替代拜访频次/级别检测\n'
                '- 模型3.2：制度内嵌名单替代附表6战略客户识别\n'
                '- 模型3.3：中标报量/签约报量标记交叉校验替代评级得分核查\n'
                '台账接入后可将上述代理检测升级为精确版本(拜访频次/级别、评级得分底线、评级调整审批记录)。'
                '当前不再阻塞维度三运行。'
            )
        break

print('Phase 7: Appendix B done')

# === SAVE ===
ts = time.strftime('%H%M%S')
DST = SRC.replace('_173519.docx', f'_v2.9_DMP驱动版_{ts}.docx')
doc.save(DST)
print(f'Saved: {DST}')
print('Done!')
