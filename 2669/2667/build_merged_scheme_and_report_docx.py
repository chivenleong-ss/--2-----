from pathlib import Path

from docx import Document


DOCS = [
    (
        Path("方案_v2.9_模型2.5签约履约偏差_附表SAP合并完整版修订.md"),
        Path("方案_v2.9_模型2.5签约履约偏差_附表SAP合并完整版修订.docx"),
    ),
    (
        Path("附件5：数字化审计模型-数字化营销审计智能分析与穿透预警模型6_附表SAP合并完整版修订.md"),
        Path("附件5：数字化审计模型-数字化营销审计智能分析与穿透预警模型6_附表SAP合并完整版修订.docx"),
    ),
]


def add_line(doc: Document, line: str) -> None:
    text = line.rstrip()
    if not text:
        doc.add_paragraph("")
        return
    if text.startswith("# "):
        doc.add_heading(text[2:].strip(), level=0)
        return
    if text.startswith("## "):
        doc.add_heading(text[3:].strip(), level=1)
        return
    if text.startswith("### "):
        doc.add_heading(text[4:].strip(), level=2)
        return
    doc.add_paragraph(text)


def convert_one(src: Path, dst: Path) -> None:
    doc = Document()
    for line in src.read_text(encoding="utf-8").splitlines():
        add_line(doc, line)
    doc.save(dst)
    print(dst)


def main() -> None:
    for src, dst in DOCS:
        convert_one(src, dst)


if __name__ == "__main__":
    main()
