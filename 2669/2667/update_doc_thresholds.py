"""
Final update to the implementation plan: add threshold corrections from 手册0520 policy verification.
"""
import sys, time
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Use the most recent final version
SRC = r'C:\Users\sasa\Desktop\模型建设\模型2：市场营销\【完善版修订】全面数字化营销审计系统实施方案_v2.9_final.docx'

doc = Document(SRC)
body = doc.element.body

def add_heading_para(text, level=1):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), f'Heading{level}')
    pPr.append(pStyle)
    p.append(pPr)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    body.append(p)

def add_body_para(text):
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    body.append(p)

# ====================================================================
# 1. Update title & revision
# ====================================================================
print('1. Updating title and revision...')
doc.paragraphs[0].text = '全面数字化营销审计系统端到端落地实施方案（完善版修订稿 v2.9）'
doc.paragraphs[1].text = '**修订日期**：2026-06-02'

# Update P6 to add threshold corrections
old_p6 = doc.paragraphs[6].text
if '制度阈值校验' not in old_p6:
    doc.paragraphs[6].text = (
        'v2.9核心更新：（1）DMP四节点时间线校验替代OA审批流（交标→中标→签约报量→签约），'
        '模型1.4新增4项时序倒置检测；（2）模型2.1新增「是否付款条件不达标」制度交叉校验；'
        '（3）签约报量（四局）170字段导出适配；（4）三证确认DMP已含；'
        '（5）offline_fields从40字段扩展至63字段；'
        '（6）代码审查修复：model_2.2 zombie_days、report摘要key；'
        '（7）OA策略调整为"DMP初筛→疑点定向OA核查"二级模式；'
        '（8）前置过滤层及模型2.2「停工/退场」统一扩展为「停工/退场/停缓建」；'
        '（9）🔴 制度阈值全面校验：对照手册0520附件4/5逐项修正8类偏差（非现金支付12%→15%、'
        '利润率底线8%→0%、履约监控阈值等），详见新增第八章。'
    )

# ====================================================================
# 2. Update P13 (DMP 170 fields) if still says 173
# ====================================================================
if '173' in doc.paragraphs[13].text:
    doc.paragraphs[13].text = doc.paragraphs[13].text.replace('173', '170')

# ====================================================================
# 3. Update section 3.1 description to note corrections
# ====================================================================
for i, p in enumerate(doc.paragraphs):
    if '3.1 制度阈值' in p.text and p.style.name.startswith('Heading'):
        # Find the paragraph after the JSON block that describes it
        for j in range(i, min(i+100, len(doc.paragraphs))):
            if '制度切换' in doc.paragraphs[j].text and doc.paragraphs[j].style.name != 'Heading':
                # This is inside the JSON-like display. Find the end.
                pass
        print(f'  Section 3.1 starts at P{i}')
        break

# ====================================================================
# 4. Add Chapter 8: 制度阈值校验报告 (after Chapter 7)
# ====================================================================
print('2. Adding Chapter 8: 制度阈值校验报告...')

# Find the last appendix heading before 附录A
ch7_end = None
for i, p in enumerate(doc.paragraphs):
    if '附录A' in p.text and p.style.name.startswith('Heading'):
        ch7_end = i
        break

if ch7_end:
    # Insert before 附录A
    # We add paragraphs to the body before the 附录A element
    ref_element = doc.paragraphs[ch7_end]._element

    def insert_para_before(ref, text, level=None):
        p = OxmlElement('w:p')
        if level:
            pPr = OxmlElement('w:pPr')
            pStyle = OxmlElement('w:pStyle')
            pStyle.set(qn('w:val'), f'Heading{level}')
            pPr.append(pStyle)
            p.append(pPr)
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        p.append(r)
        ref.addprevious(p)

    # Add Chapter 8 content
    insert_para_before(ref_element, '', None)
    insert_para_before(ref_element, '', None)
    insert_para_before(ref_element, '第八章 制度阈值校验报告（v2.9 对照手册0520附件4/5）', 1)
    insert_para_before(ref_element, '', None)
    insert_para_before(ref_element,
        'v2.9版本对照《中国建筑第四工程局有限公司市场营销管理手册0520》（2026年5月修订版）'
        '附件4（严禁投标事项清单）和附件5（限制投标事项清单），对config/rules.json中所有'
        '制度阈值进行了逐项校验。以下为发现并修正的8类重大偏差：', None)
    insert_para_before(ref_element, '', None)

    insert_para_before(ref_element, '8.1 已修正的错误阈值', 2)
    insert_para_before(ref_element, '', None)

    corrections = [
        ('1. 地产类非现金支付上限：12% → 15%',
         '手册0520附件4第11.5条明确：地产类项目非现金支付比例超过合同额15%即触发严禁投标。'
         '旧版rules.json误设为12%，已修正为15%。影响模型2.1的地产类非现金支付超标检测。'),
        ('2. 非地产类非现金支付上限：30% → 15%',
         '手册0520附件4第9.3条：非地产类项目非现金支付比例超过合同额15%为组合判定条件之一。'
         '旧版rules.json沿用2020版旧标准的30%，严重偏宽。已修正为15%。'),
        ('3. 地产类月进度付款比例下限（新制度）：统一0.75 → 橙档0.75 / 黄绿档0.70',
         '手册0520附件4第11.1条区分了橙档（<75%→严禁）和黄绿档（<70%→严禁）。'
         '旧版rules.json手册0520部分未区分监管档次，统一使用0.75。已拆分为两个独立阈值。'),
        ('4. 非地产类预期净利润率下限：8% → 移除（沿用全项目通用<0%严禁线）',
         '手册0520附件4第6条规定"承接效益率低于0的项目"为严禁投标，这是全项目通用的唯一利润率严禁线。'
         '手册0520中不存在8%这一阈值。旧版rules.json的8%可能来源于被废止的2020版制度或其他内部标准。'
         '已移除该阈值，改为统一使用承接效益率严禁投标上限（0%）。'),
        ('5. 地产类非橙档预期净利润率下限：8% → 移除',
         '同上。手册0520对非橙档地产类项目仅规定了<0%为严禁投标（附件4第6条），'
         '无额外的8%下限。橙档地产额外有0-5%的限制投标区间（附件5第9条）。'),
        ('6. 基础设施预期施工利润率下限：10% → 移除',
         '手册0520中未对基础设施类项目设置单独的最低利润率标准。'
         '全项目通用的<0%严禁线（附件4第6条）适用于所有项目类型。'
         '垫资≥5000万的非地产房建/专业工程有4%/5%的要求（附件5第2条），但这是垫资触发条件。'),
        ('7. 风险分级履约监控阈值：全部修正',
         '旧版rules.json中的履约监控数值与手册0520附件4第10条（非地产类）和第13条（地产类）'
         '存在系统性偏差。例如：橙档已结算未收款占比：旧值2%→正确值3%；'
         '黄档已结算未收款占比：旧值3%→正确值5%；绿档已结算未收款占比：旧值5%→正确值8%。'
         '非现金占比的偏差尤为严重：橙档/黄档旧值8%/10%→正确值均为0%。已全部按原文修正。'),
        ('8. 现金履约保证金阈值硬编码 → 配置化',
         '模型2.3中现金履约保证金的严重等级判定使用硬编码100万元阈值。'
         '已改为从config/rules.json读取「现金履约保证金_二级单位审批上限_万元」（手册0520附件5第3条）。'),
    ]

    for title, desc in corrections:
        insert_para_before(ref_element, f'**{title}**', None)
        insert_para_before(ref_element, desc, None)
        insert_para_before(ref_element, '', None)

    insert_para_before(ref_element, '8.2 修改文件清单', 2)
    insert_para_before(ref_element, '', None)

    files_changed = [
        'config/rules.json — 全面重写付款条件底线、盈利底线、风险分级履约监控三个配置块，新增非地产组合判定条件、地产类完工/竣工/结算比例阈值',
        'models/dim2/model_2_1_risk.py — 新制度月进度付款区分橙档/黄绿档；非现金支付上限修正；利润率全项目通用<0%严禁线；非地产非现金支付上限修正',
        'models/dim2/model_2_2_profit.py — 新制度利润率底线统一改为承接效益率严禁投标上限(0.0)；移除不存在的8%/10%阈值引用',
        'models/dim2/model_2_3_capital.py — 现金履约保证金阈值从硬编码改为config读取',
    ]

    for f in files_changed:
        insert_para_before(ref_element, f'• {f}', None)

    insert_para_before(ref_element, '', None)
    insert_para_before(ref_element, '8.3 新增至config但代码尚未消费的阈值', 2)
    insert_para_before(ref_element, '', None)

    pending = [
        '非地产组合判定（附件4第9条）：4项条件中≥2项触发严禁投标。当前代码每项独立判定，建议后续版本实现组合逻辑。',
        '完工/竣工/结算比例（附件4第11.2-3条）：完工累计付款<80%或竣工备案<85%、竣工结算<95%。已写入config，模型代码尚未消费。',
        '付款周期检查（附件4第11.1条、附件5第1条）：橙档≤2月/黄绿档≤3月。config已有字段，代码仅检查月进度比例未检查付款周期。',
        '垫资审批层级联动（附件5第2条）：垫资≥5000万由局董事会审批，2000万≤垫资<5000万由二级单位董事会审批。<2000万由二级单位董事长审批。config已配置但代码仅在模型2.3用作阈值参考。',
    ]

    for item in pending:
        insert_para_before(ref_element, f'• {item}', None)

    insert_para_before(ref_element, '', None)
    insert_para_before(ref_element, '8.4 确认无需修改的部分', 2)
    insert_para_before(ref_element, '', None)

    unchanged = [
        '旧制度（2020版）阈值：橙档5%/黄绿档4%、付款不达标时的差异化底线 — 与原2020版制度一致',
        '区域管理阈值：跨区域合同额门槛5亿（手册0520附件5）、历史10亿、安装公司1.5亿',
        '客户管理/经验预警：评级合格线60分、战略客户拜访年均4次、客户集中度60% — 内部经验值，非手册0520直接规定',
        '十五五战略规划基准 — 来源于十五五规划征求意见稿，非手册0520',
        '地产类任意垫资即严禁（附件4第11.4条）— 代码中「是否垫资=是」的判定逻辑正确',
        '地产类任意现金履约保证金即严禁（附件4第11.6条）— 代码中检查履约担保方式含"现金"的逻辑正确',
    ]

    for item in unchanged:
        insert_para_before(ref_element, f'• {item}', None)

    print('  Chapter 8 added before 附录A')

# ====================================================================
# 5. Save
# ====================================================================
ts = time.strftime('%H%M%S')
DST = SRC.replace('_final.docx', f'_v2.9_complete.docx')
doc.save(DST)
print(f'\nSaved: {DST}')
print(f'Filename: {DST.split(chr(92))[-1]}')
print('Done!')
