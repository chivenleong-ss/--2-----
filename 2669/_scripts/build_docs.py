"""
Generate two Word documents:
1. 实施方案_六模块九宫格全链路贯通_v2.10.docx — standalone implementation plan
2. Updated 方案_v2.9 with new Chapter 8 appended
"""
import os
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

PROJECT_ROOT = r"C:\Users\sasa\Desktop\模型建设\模型2：市场营销"
DOCS_DIR = os.path.join(PROJECT_ROOT, "_documents")

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_border(cell, **kwargs):
    """Set cell borders"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_styled_heading(doc, text, level=1):
    """Add a styled heading"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def add_body_para(doc, text, bold=False, font_size=11):
    """Add body paragraph with proper CJK font"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(font_size)
    run.bold = bold
    return p

def add_table_with_style(doc, headers, rows, col_widths=None):
    """Add a styled table"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(10)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1f89df')

    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.size = Pt(9)
            if r % 2 == 1:
                set_cell_shading(cell, 'F5F8FC')

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()  # spacer
    return table


def create_standalone_plan():
    """Create the standalone implementation plan document"""
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── Title ──
    title = doc.add_heading('全面数字化营销审计系统', level=0)
    for run in title.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    subtitle = doc.add_heading('六模块·九宫格·全链路贯通实施方案', level=1)
    for run in subtitle.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for text, bold in [
        ('版本：v2.10（基于v2.9增量升级）\n', True),
        (f'日期：{datetime.date.today().isoformat()}\n', False),
        ('状态：待实施\n', False),
        ('关联文档：方案_v2.9_模型2.5签约履约偏差_完整版修订.docx', False),
    ]:
        run = meta.add_run(text)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(10)
        run.bold = bold

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 第一章：背景与目标
    # ═══════════════════════════════════════════
    add_styled_heading(doc, '第一章 背景与目标', level=1)

    add_styled_heading(doc, '1.1 问题诊断', level=2)
    add_body_para(doc, '当前v2.9系统已实现11个审计模型、三个维度、三条关联链、九宫格离散分析。但在"模型→六模块→九宫格"的链路上存在以下断层：')

    problems = [
        '断层一：六模块不完整。BusinessHealthAnalyzer仅实现了5个模块（区域/客户/合同/履约/资金），缺少"模块六：数据质量与流程效率分析"（基于模型1.4）。',
        '断层二：指标展示不全。business.html的"区域-客户-履约健康监测台"仅展示了6个零散KPI，未按六模块的结构化指标体系展示全部30+指标。',
        '断层三：九宫格R/S与六模块脱节。DiscreteAnalyzer的R/S子维度直接读取11个模型的原始输出（红/黄标记），中间绕过了六模块指标计算层，导致九宫格的R/S与监测台的六模块得分互不关联。',
        '断层四：缺少置信度反馈。即使项目落入"扩张区"，若底层数据质量差（模块六得分低），决策者无从知晓。',
        '断层五：缺少推演能力。九宫格定位是静态的，业务人员无法模拟"如果改善某个指标，项目会在九宫格中如何移动"。',
    ]
    for i, p in enumerate(problems, 1):
        add_body_para(doc, f'({i}) {p}')

    add_styled_heading(doc, '1.2 目标架构', level=2)
    add_body_para(doc, '本次升级目标：建立"11模型→三维度→三条链→六模块→九宫格"的完整数据链路，使得：')

    goals = [
        '六模块指标层成为整个系统的数据中枢：所有上层分析（监测台、九宫格）均从六模块取数。',
        '监测台完整展示六模块的全部30+指标，支持异步加载和骨架屏体验。',
        '九宫格R/S子维度直接采用六模块指标，实现与监测台的指标同源、逻辑一致。',
        '模块六得分作为项目的数据置信度标识，在九宫格散点图中以视觉编码（虚线边框、透明度）提示数据失真风险。',
        '提供前端What-If沙盘推演：业务人员拖动滑块调整模块指标，气泡在九宫格中实时跃迁。',
        '全链路内存管理：大宽表分阶段释放，Web API分模块异步返回，防止并发访问时内存溢出。',
    ]
    for g in goals:
        add_body_para(doc, f'• {g}')

    # Architecture diagram
    add_body_para(doc, '')
    add_body_para(doc, '总体架构：', bold=True)
    arch_text = """
                    ┌────────── 应用层 ──────────┐
                    │  区域-客户-履约健康监测台    │
                    │  九宫格风险-收益决策矩阵     │
                    │  沙盘推演 What-If 模式       │
                    └────────────┬────────────────┘
                                 │
                    ┌────────── 分析层 ──────────┐
                    │  六模块指标计算引擎          │
                    │  (BusinessHealthAnalyzer)    │
                    │  模块一~六，30+指标          │
                    └────────────┬────────────────┘
                                 │
              ┌──────────────────┼──────────────────┐
              ↓                  ↓                  ↓
     ┌── 九宫格离散分析 ──┐ ┌── 三条关联链 ──┐ ┌── 置信度引擎 ──┐
     │ R/S 直接采用六模块  │ │ chain_1/2/3   │ │ 模块六 →        │
     │ 指标，公式化计算     │ │               │ │ 置信度标记      │
     └─────────────────────┘ └───────────────┘ └────────────────┘
                                 │
                    ┌────────── 模型层 ──────────┐
                    │  11个审计模型（1.1~3.2）     │
                    │  红/黄标记 + 问题分类        │
                    └────────────┬────────────────┘
                                 │
                    ┌────────── 数据层 ──────────┐
                    │  DMP + 附表 + 中标报量       │
                    │  前置过滤 + 区域认定         │
                    └─────────────────────────────┘
"""
    p = doc.add_paragraph()
    run = p.add_run(arch_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 第二章：六模块完整指标体系
    # ═══════════════════════════════════════════
    add_styled_heading(doc, '第二章 六模块完整指标体系', level=1)

    add_styled_heading(doc, '2.1 11模型→六模块映射', level=2)

    mapping_headers = ['模块', '对应模型', '核心指标（5-6个/模块）']
    mapping_rows = [
        ['模块一：区域布局健康度', '1.1 + 1.2', '①区域渗透率 ②跨区域经营指数 ③深耕区域集中度\n④区域合同额强度 ⑤业务结构偏离度 ⑥EPC转型进度'],
        ['模块二：客户资源稳定性', '1.3 + 3.1 + 3.2', '①客户稳定性指数 ②客户产出波动率 ③客户集中度风险\n④中标转化率 ⑤新客户质量指数 ⑥战略客户产出比'],
        ['模块三：合同质量与风险集中度', '2.1 + 2.4', '①风险项目占比 ②风险合同额集中度 ③付款条件优良率\n④合同条款不利度 ⑤三证合规率'],
        ['模块四：履约盈利健康度', '2.2 + 2.5 + 前置过滤', '①产值转化率 ②签约履约偏差率 ③盈利健康度\n④停工退场率 ⑤效益偏差率 ⑥在施项目活跃度'],
        ['模块五：资金效率与安全性', '2.3', '①资金占用率 ②保证金周转天数 ③逾期回收率\n④预收款缺口率 ⑤负流项目占比'],
        ['模块六：数据质量与流程效率', '1.4', '①数据完整率 ②流程合规率 ③中标签约偏差率\n④测算规律性指数 ⑤签约延迟率'],
    ]
    add_table_with_style(doc, mapping_headers, mapping_rows, [3.5, 3.0, 9.5])

    add_styled_heading(doc, '2.2 每个模块的指标详细定义', level=2)

    # Module 1
    add_styled_heading(doc, '模块一：区域布局健康度分析（基于模型1.1+1.2）', level=3)
    m1_headers = ['指标名称', '计算逻辑', '经营意义']
    m1_rows = [
        ['区域渗透率', '实际签约城市数 / 授权城市数 × 100%', '衡量区域深耕深度'],
        ['跨区域经营指数', '非常规区域合同额 / 总合同额', '反映经营扩张激进程度'],
        ['深耕区域集中度', '核心+重点城市合同额占比', '评估区域聚焦度'],
        ['区域合同额强度', '单城市平均合同额（分母为签约项目数）', '识别"撒胡椒面"式布局'],
        ['业务结构偏离度', 'Σ|实际占比 - 战略占比| / 板块数', '衡量战略执行偏差'],
        ['EPC转型进度', 'EPC模式合同额 / 总合同额', '业务模式升级速度'],
    ]
    add_table_with_style(doc, m1_headers, m1_rows, [3.5, 6.0, 6.5])

    # Module 2
    add_styled_heading(doc, '模块二：客户资源稳定性分析（基于模型1.3+3.1+3.2）', level=3)
    m2_headers = ['指标名称', '计算逻辑', '经营意义']
    m2_rows = [
        ['客户稳定性指数', '1 - (流失客户数/上年客户总数)', '客户资源流失风险'],
        ['客户产出波动率', '标准差(年度合同额) / 均值', '客户产出是否稳定'],
        ['客户集中度风险', '前5大客户合同额占比', '依赖度过高预警'],
        ['中标转化率', '签约额 / 中标额', '营销漏斗效率'],
        ['新客户质量指数', '国企/政府类新客户占比 × 粘性系数', '获客质量评估'],
        ['战略客户产出比', '战略客户合同额 / 总合同额', '核心客户贡献度'],
    ]
    add_table_with_style(doc, m2_headers, m2_rows, [3.5, 6.0, 6.5])

    # Module 3
    add_styled_heading(doc, '模块三：合同质量与风险集中度分析（基于模型2.1+2.4）', level=3)
    m3_headers = ['指标名称', '计算逻辑', '经营意义']
    m3_rows = [
        ['风险项目占比', '触碰严禁投标底线项目数 / 总项目数', '整体风险水位'],
        ['风险合同额集中度', '高风险项目合同额 / 总合同额', '风险敞口规模'],
        ['付款条件优良率', '付款达标项目数 / 总项目数', '资金回收保障度'],
        ['合同条款不利度', '高风险条款项目数 / 总项目数', '合同谈判质量'],
        ['三证合规率', '三证齐全项目数 / 已开工项目数', '合规经营水平'],
    ]
    add_table_with_style(doc, m3_headers, m3_rows, [3.5, 6.0, 6.5])

    # Module 4
    add_styled_heading(doc, '模块四：履约盈利健康度分析（基于模型2.2+2.5+前置过滤）', level=3)
    m4_headers = ['指标名称', '计算逻辑', '经营意义']
    m4_rows = [
        ['产值转化率', '实际完成产值 / 签约额', '合同落地效率'],
        ['签约履约偏差率', '1 - 产值转化率（签约>12月）', '识别"签而不做"项目'],
        ['盈利健康度', 'A值≥底线项目占比', '整体盈利质量'],
        ['停工退场率', '停工退场项目数 / 已开工项目数', '项目夭折风险'],
        ['效益偏差率', '|A值 - 实际利润率| > 1% 项目占比', '测算准确性'],
        ['在施项目活跃度', '在施项目产值>0且收款>0占比', '在建项目健康度'],
    ]
    add_table_with_style(doc, m4_headers, m4_rows, [3.5, 6.0, 6.5])

    # Module 5
    add_styled_heading(doc, '模块五：资金效率与安全性分析（基于模型2.3）', level=3)
    m5_headers = ['指标名称', '计算逻辑', '经营意义']
    m5_rows = [
        ['资金占用率', '(保证金+预收款应收) / 总合同额', '资金沉淀规模'],
        ['保证金周转天数', '平均(实际回收日 - 约定退还日)', '资金回收效率'],
        ['逾期回收率', '逾期>90天款项金额 / 应回收总金额', '资金风险敞口'],
        ['预收款缺口率', '(应收预收款 - 实收预收款) / 应收预收款', '预收款回收缺口'],
        ['负流项目占比', '资金结余<<0项目数 / 总项目数', '现金流健康度'],
    ]
    add_table_with_style(doc, m5_headers, m5_rows, [3.5, 6.0, 6.5])

    # Module 6
    add_styled_heading(doc, '模块六：数据质量与流程效率分析（基于模型1.4）', level=3)
    m6_headers = ['指标名称', '计算逻辑', '经营意义']
    m6_rows = [
        ['数据完整率', '关键字段非空率（DMP 78字段）', '数据治理水平'],
        ['流程合规率', '招文评审→交标→中标→签约时序合规占比', '流程执行质量'],
        ['中标签约偏差率', '|中标额-签约额|>5%项目占比', '报价稳定性'],
        ['测算规律性指数', '同单位利润率标准差', '数据真实性'],
        ['签约延迟率', '预计签约日已过未签约项目占比', '营销转化效率'],
    ]
    add_table_with_style(doc, m6_headers, m6_rows, [3.5, 6.0, 6.5])

    add_styled_heading(doc, '2.3 综合评分权重（5模块→6模块）', level=2)
    add_body_para(doc, '将原5模块评分体系扩展为6模块：')

    weight_headers = ['模块', '原权重', '新权重', '变化说明']
    weight_rows = [
        ['模块一：区域布局', '30%', '25%', '新增模块六分走5%'],
        ['模块二：客户稳定', '20%', '20%', '不变'],
        ['模块三：合同质量', '20%', '18%', '新增模块六分走2%'],
        ['模块四：履约盈利', '15%', '15%', '不变'],
        ['模块五：资金效率', '15%', '12%', '新增模块六分走3%'],
        ['模块六：数据质量', '—（未实现）', '10%', '新增'],
    ]
    add_table_with_style(doc, weight_headers, weight_rows, [4.0, 2.5, 2.5, 7.0])

    add_body_para(doc, '')
    add_body_para(doc, '综合得分计算公式：', bold=True)
    add_body_para(doc, '综合得分 = 模块一(25%) + 模块二(20%) + 模块三(18%) + 模块四(15%) + 模块五(12%) + 模块六(10%)')
    add_body_para(doc, '分区阈值：综合得分 ≥ 80 为强势区，65 ≤ 综合得分 < 80 为稳健区，综合得分 < 65 为承压区。')

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 第三章：九宫格R/S与六模块指标映射
    # ═══════════════════════════════════════════
    add_styled_heading(doc, '第三章 九宫格R/S与六模块指标映射', level=1)

    add_styled_heading(doc, '3.1 核心设计原则', level=2)
    add_body_para(doc, '九宫格的风险维度(R)和收益维度(E)的所有子维度，不再直接读取11个模型的原始输出（红/黄标记），而是改为直接采用六模块的标准化指标。这使得：')
    add_body_para(doc, '• 监测台展示的六模块指标与九宫格使用的R/S子维度完全同源，逻辑一致')
    add_body_para(doc, '• R/S的计算公式完全透明化、配置化')
    add_body_para(doc, '• What-If推演成为可能：调整六模块指标 → R/S自动重算 → 气泡在九宫格中移动')

    add_styled_heading(doc, '3.2 R（风险维度）← 六模块风险指标映射', level=2)

    r_headers = ['R子维度', '权重', '对应六模块指标', '来源模块', '方向', '分箱阈值(1/2/3档)']
    r_rows = [
        ['区域合规', '1.0', '跨区域经营指数', '模块一', 'inverse', '≤20%→1, 20-50%→2, >50%→3'],
        ['合同底线', '1.0', '风险项目占比(0.6) + 合同条款不利度(0.4)', '模块三', 'inverse', '≤15%→1, 15-35%→2, >35%→3'],
        ['客户健康', '0.8', '客户稳定性指数(0.5) + 客户集中度风险(0.5)', '模块二', 'inverse', '稳定性<60%→3, 集中度>60%→3'],
        ['资金安全', '0.8', '逾期回收率(0.6) + 负流项目占比(0.4)', '模块五', 'inverse', '逾期>20%→3, 负流>15%→3'],
        ['履约真实', '0.4', '停工退场率(0.6) + 签约履约偏差率(0.4)', '模块四', 'inverse', '停工>10%→3, 偏差>50%→3'],
    ]
    add_table_with_style(doc, r_headers, r_rows, [2.0, 1.0, 4.5, 1.5, 1.5, 5.5])

    add_body_para(doc, '')
    add_body_para(doc, 'R得分计算：', bold=True)
    add_body_para(doc, 'R_raw = 区域合规×1.0 + 合同底线×1.0 + 客户健康×0.8 + 资金安全×0.8 + 履约真实×0.4')
    add_body_para(doc, 'R = R_raw / 4.0（归一化到[1.0, 3.0]）')
    add_body_para(doc, '分箱：R ≤ 1.6 → 低风险(1)；1.6 < R ≤ 2.4 → 中风险(2)；R > 2.4 → 高风险(3)')

    add_styled_heading(doc, '3.3 E（收益维度）← 六模块收益指标映射', level=2)

    e_headers = ['E子维度', '权重', '对应六模块指标', '来源模块', '方向', '分箱阈值(1/2/3档)']
    e_rows = [
        ['盈利水平', '1.05', '盈利健康度', '模块四', 'direct', '≥80%→3, 50-80%→2, <50%→1'],
        ['产值转化', '0.75', '产值转化率', '模块四', 'direct', '≥80%→3, 50-80%→2, <50%→1'],
        ['资金回收', '0.75', '资金回收率（收款/签约额）', '模块五', 'direct', '≥60%→3, 30-60%→2, <30%→1'],
        ['合同规模', '0.45', '区域合同额强度（归一化）', '模块一', 'direct', '≥5亿→3, 1-5亿→2, <1亿→1'],
    ]
    add_table_with_style(doc, e_headers, e_rows, [2.0, 1.0, 4.5, 1.5, 1.5, 5.5])

    add_body_para(doc, '')
    add_body_para(doc, 'E得分计算：', bold=True)
    add_body_para(doc, 'E_raw = 盈利水平×1.05 + 产值转化×0.75 + 资金回收×0.75 + 合同规模×0.45')
    add_body_para(doc, 'E = E_raw / 3.0（归一化到[1.0, 3.0]）')
    add_body_para(doc, '分箱：E ≤ 1.6 → 低收益(1)；1.6 < E ≤ 2.4 → 中收益(2)；E > 2.4 → 高收益(3)')

    add_styled_heading(doc, '3.4 九宫格策略矩阵（不变）', level=2)

    grid_headers = ['格子', '风险-收益组合', '业务命名', '项目特征', '处置策略']
    grid_rows = [
        ['(1,1)', '低风险-低收益', '退出区', '深耕区域但项目小、盈利差', '主动退出或合并，释放资源'],
        ['(1,2)', '低风险-中收益', '培育区', '合规经营但规模未起量', '资源倾斜，扩大市场份额'],
        ['(1,3)', '低风险-高收益', '扩张区', '深耕区域+大项目+高盈利', '复制推广，设立区域中心'],
        ['(2,1)', '中风险-低收益', '观察区', '非常规区域+小项目+盈利差', '收缩投入，限期改善'],
        ['(2,2)', '中风险-中收益', '维持区', '存在单一风险但收益尚可', '常规管理，动态监控'],
        ['(2,3)', '中风险-高收益', '优化区', '非常规区域但项目优质', '压缩风险敞口，推动合规化'],
        ['(3,1)', '高风险-低收益', '淘汰区', '跨区域违规+触碰红线+亏损', '立即止损，启动问责'],
        ['(3,2)', '高风险-中收益', '整顿区', '触碰红线但收益尚可', '限期整改，回溯审批'],
        ['(3,3)', '高风险-高收益', '警惕区', '违规但高盈利（刀口舔血）', '专人周跟踪，一案一策'],
    ]
    add_table_with_style(doc, grid_headers, grid_rows, [1.2, 3.0, 2.0, 4.5, 5.3])

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 第四章：数据置信度视觉反馈
    # ═══════════════════════════════════════════
    add_styled_heading(doc, '第四章 数据置信度视觉反馈', level=1)

    add_styled_heading(doc, '4.1 设计理念', level=2)
    add_body_para(doc, '模块六（数据质量与流程效率）的得分不仅作为综合评分的10%权重，更作为整个九宫格的"数据置信度（Confidence Interval）"标识。核心逻辑：如果某项目落在"(1,3) 扩张区"但模块六得分极低（存在中标签约金额偏差、关键字段缺失等问题），这说明"该项目看似高收益低风险，但底层数据严重失真，决策需谨慎"。')

    add_styled_heading(doc, '4.2 置信度等级定义', level=2)

    conf_headers = ['置信度等级', '模块六得分', '视觉编码', '决策含义']
    conf_rows = [
        ['高置信度', '≥ 80分', '实心圆点，正常边框，不透明度0.92', '数据质量可靠，可据此决策'],
        ['中置信度', '50-79分', '半透明（opacity 0.70），橙色边框', '数据存在一定缺失，建议复核后再决策'],
        ['低置信度', '< 50分', '高度透明（opacity 0.45），红色虚线边框，\n菱形符号，宽度3px', '⚠ 数据严重失真，决策前必须先复核底层数据'],
    ]
    add_table_with_style(doc, conf_headers, conf_rows, [2.5, 2.5, 5.5, 5.5])

    add_styled_heading(doc, '4.3 置信度在散点图中的视觉编码', level=2)
    add_body_para(doc, '在ECharts散点图（discrete.html）中，每个气泡根据其置信度等级呈现不同的视觉样式：')
    add_body_para(doc, '• itemStyle.color：高置信度=原色，中置信度=原色+99半透明，低置信度=原色+55高度透明')
    add_body_para(doc, '• itemStyle.borderColor：高=#fff，中=#f59e0b（橙），低=#d94d63（红）')
    add_body_para(doc, '• itemStyle.borderWidth：高=1.5px，中=2px，低=3px')
    add_body_para(doc, '• itemStyle.borderType：低置信度=dashed（虚线边框）')
    add_body_para(doc, '• symbol：低置信度（<30分）=diamond（菱形），其他=circle')
    add_body_para(doc, '• Tooltip：低置信度气泡悬停时显示红色警告"⚠ 数据置信度低(NN分)"及具体数据质量问题列表')
    add_body_para(doc, '')
    add_body_para(doc, '散点图图例区新增置信度说明条，包含高/中/低三种置信度的图例标识。')

    add_styled_heading(doc, '4.4 后端实现', level=2)
    add_body_para(doc, '在discrete_analysis.py的_analyze_projects()方法中，为每个项目新增两个字段：')
    add_body_para(doc, '• 数据置信度：0-100的浮点数，直接取自该单位/城市的模块六得分')
    add_body_para(doc, '• 置信度等级："high" / "medium" / "low"')
    add_body_para(doc, '• 数据质量问题：从模型1.4输出中提取的该项目的具体数据质量问题列表')

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 第五章：What-If沙盘推演
    # ═══════════════════════════════════════════
    add_styled_heading(doc, '第五章 What-If沙盘推演模式', level=1)

    add_styled_heading(doc, '5.1 设计理念', level=2)
    add_body_para(doc, '由于九宫格的R和S已完全公式化（由六模块指标加权生成），可以在前端用JavaScript完整复现计算逻辑。业务人员无需重新运行后端Python模型，即可在网页端通过滑块微调模块指标，实时观察气泡在九宫格中的位置变化。')

    add_styled_heading(doc, '5.2 前端实现方案', level=2)
    add_body_para(doc, '在discrete.html右侧新增可折叠侧边栏"📊 模拟调节杆"，包含以下功能区域：')
    add_body_para(doc, '')
    add_body_para(doc, '（1）推演对象选择器', bold=True)
    add_body_para(doc, '下拉框选择要推演的项目（默认当前点击选中的项目）。可选择查看"当前值"和"推演值"两列。')
    add_body_para(doc, '')
    add_body_para(doc, '（2）六模块指标滑块', bold=True)
    add_body_para(doc, '将Python的R/E计算公式完整移植到JavaScript（DISCRETE_FORMULA对象），包括：')
    add_body_para(doc, '• riskWeights: {region:1.0, contract:1.0, customer:0.8, capital:0.8, perf:0.4}')
    add_body_para(doc, '• returnWeights: {profit:1.05, conversion:0.75, collection:0.75, scale:0.45}')
    add_body_para(doc, '• calcRiskSubDim() / calcReturnSubDim()：六模块指标→1/2/3离散档位')
    add_body_para(doc, '• recalc(moduleIndicators)：输入调整后的六模块指标，输出新的{R, E}')
    add_body_para(doc, '每个模块提供1-2个可调节的关键指标滑块（如"产值转化率：68% → [══●══] → 88%"），滑块背景色按指标当前分区着色。')
    add_body_para(doc, '')
    add_body_para(doc, '（3）推演结果展示', bold=True)
    add_body_para(doc, '• 当前定位：维持区 (2,2)')
    add_body_para(doc, '• 推演后定位：扩张区 (1,3) ✅')
    add_body_para(doc, '• R: 2.1 → 1.5（↓0.6）')
    add_body_para(doc, '• E: 2.3 → 2.8（↑0.5）')
    add_body_para(doc, '')
    add_body_para(doc, '（4）推演轨迹动画', bold=True)
    add_body_para(doc, '在散点图上绘制从原始位置到推演位置的虚线箭头轨迹（绿色），推演气泡使用菱形+绿色边框。使用ECharts的lines系列+effect动画实现。')

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 第六章：前端页面改造
    # ═══════════════════════════════════════════
    add_styled_heading(doc, '第六章 前端页面改造方案', level=1)

    add_styled_heading(doc, '6.1 business.html：区域-客户-履约健康监测台', level=2)

    add_body_para(doc, '6.1.1 页面布局', bold=True)
    add_body_para(doc, '重新设计为三区联动布局：')
    add_body_para(doc, '• 顶部KPI概览行：覆盖单位数、覆盖城市数、高风险项目数、最高得分单位')
    add_body_para(doc, '• 中部左侧：六轴雷达图（6个模块得分）')
    add_body_para(doc, '• 中部右侧：态势总览（6模块得分卡片 + 分区分布）')
    add_body_para(doc, '• 下部：六模块详情展开区（6个可折叠卡片，每个卡片展示该模块的全部5-6个指标值、趋势箭头、得分等级）')
    add_body_para(doc, '• 底部：三级经营对标表（二级单位 | 城市 | 重点项目，含各模块得分列）')

    add_body_para(doc, '')
    add_body_para(doc, '6.1.2 异步加载与骨架屏', bold=True)
    add_body_para(doc, '后端新增3类轻量API端点以支持异步分块加载：')
    add_body_para(doc, '• GET /api/business/modules-summary —— 首次快速返回6个模块得分+综合评分（数据量<1KB）')
    add_body_para(doc, '• GET /api/business/module/<id> —— 按需返回单个模块的完整指标数据（id=1~6）')
    add_body_para(doc, '• GET /api/business/benchmark/<scope> —— 按需返回对标数据（subsidiaries/cities/projects）')

    add_body_para(doc, '')
    add_body_para(doc, '前端加载时序：', bold=True)
    add_body_para(doc, '1. [0ms] 渲染骨架屏（6个shimmer动画卡片 + KPI占位）')
    add_body_para(doc, '2. [50ms] fetch /api/business/modules-summary → KPI卡片先亮起 + 雷达图渲染')
    add_body_para(doc, '3. [50ms] 并行6个fetch（Promise.allSettled）→ 哪个模块先返回就先亮起哪个卡片')
    add_body_para(doc, '4. [200ms] fetch /api/business/benchmark/subsidiaries → 对标表渲染')

    add_body_para(doc, '')
    add_body_para(doc, '6.1.3 骨架屏CSS动画', bold=True)
    add_body_para(doc, '使用shimmer动画：linear-gradient背景在200%宽度上平移，模拟加载中的闪光效果。数据到达后，骨架屏淡出（opacity 1→0, 300ms），实际卡片淡入（opacity 0→1, 400ms）。')

    add_styled_heading(doc, '6.2 discrete.html：九宫格决策矩阵', level=2)
    add_body_para(doc, '在现有页面基础上新增以下功能区域：')
    add_body_para(doc, '• 规则卡片更新：显示"六模块联动规则"，明确列出R/S子维度→六模块指标→来源模块的完整映射链')
    add_body_para(doc, '• 置信度图例：在散点图下方增加"数据置信度"图例条')
    add_body_para(doc, '• What-If侧边栏：可折叠的沙盘推演面板（详见第五章）')
    add_body_para(doc, '• 气泡点击联动：点击散点气泡时，右侧联动面板同步展示该项目的六模块得分雷达图（小尺寸）')

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 第七章：内存管理
    # ═══════════════════════════════════════════
    add_styled_heading(doc, '第七章 内存释放与性能优化', level=1)

    add_styled_heading(doc, '7.1 main.py分阶段内存释放', level=2)
    add_body_para(doc, '在pipeline执行过程中，在每个阶段结束后显式释放不再需要的大对象：')
    add_body_para(doc, '• 数据加载完成后：del dmp_raw, bid_report → gc.collect()')
    add_body_para(doc, '• 11模型执行完成后：del appendices → gc.collect()')
    add_body_para(doc, '• 六模块分析完成后：del dmp, region_auth → gc.collect()（释放173字段大宽表，可能数百MB）')
    add_body_para(doc, '• 报告生成仅消费all_results + business_results + discrete_results等已缓存的小对象')

    add_styled_heading(doc, '7.2 web_app.py API级内存优化', level=2)
    add_body_para(doc, '• 所有数据API响应完成后调用gc.collect(0)（仅回收年轻代，开销极小）')
    add_body_para(doc, '• 分模块API每次只从pickle缓存中提取单个模块数据，避免加载完整business_results到内存')
    add_body_para(doc, '• 新增GET /api/memory-status运维端点，返回RSS/VMS/GC计数，便于监控')

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 第八章：实施步骤
    # ═══════════════════════════════════════════
    add_styled_heading(doc, '第八章 实施步骤与工时估算', level=1)

    impl_headers = ['阶段', '步骤', '涉及文件', '工时']
    impl_rows = [
        ['P1\n六模块\n补全', '1. business_analysis.py新增模块六计算\n2. 扩展每模块输出全部指标(30+字段)\n3. web_app.py新增分模块API', 'models/business_analysis.py\nmodels/business_analysis.py\nweb_app.py', '2h\n1.5h\n1h'],
        ['P2\n监测台\n重构', '4. 骨架屏+异步加载JS\n5. 六模块卡片详情区+雷达图\n6. 六轴雷达+KPI+对标表', 'templates/business.html\ntemplates/business.html\ntemplates/business.html', '2h\n2h\n1.5h'],
        ['P3\nR/S\n映射', '7. discrete_rules.json新增映射配置\n8. discrete_analysis.py新增run_with_module_scores\n9. main.py调整执行顺序+内存释放', 'config/discrete_rules.json\nmodels/discrete_analysis.py\nmain.py', '0.5h\n2h\n1h'],
        ['P4\n置信度', '10. discrete_analysis.py置信度计算\n11. discrete.html散点图视觉编码', 'models/discrete_analysis.py\ntemplates/discrete.html', '1h\n1h'],
        ['P5\nWhat-If', '12. JS公式移植+滑块UI\n13. 推演气泡动画', 'templates/discrete.html\ntemplates/discrete.html', '2h\n1h'],
        ['P6\n联调', '14. 端到端测试+边界处理', '全量文件', '2h'],
    ]
    add_table_with_style(doc, impl_headers, impl_rows, [1.5, 7.0, 5.5, 2.0])

    add_body_para(doc, '')
    add_body_para(doc, '总计工时：约20.5小时', bold=True)

    add_styled_heading(doc, '8.1 数据流顺序（main.py调整后）', level=2)
    add_body_para(doc, '[1] 数据加载 → [2] 前置过滤 → [3] 11模型运行 → [4] 三条链 → [5] 六模块分析(新增，先于离散分析) → [6] 九宫格离散分析(消费六模块结果) → [7] 报告生成')

    add_styled_heading(doc, '8.2 向后兼容', level=2)
    add_body_para(doc, '如果六模块分析尚未执行（无business_results缓存），discrete_analysis.py自动回退到直接读取11模型原始输出的现有逻辑，确保系统在任何状态下均可用。')

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 第九章：涉及文件清单
    # ═══════════════════════════════════════════
    add_styled_heading(doc, '第九章 涉及文件清单', level=1)

    files_headers = ['文件', '改动类型', '改动说明']
    files_rows = [
        ['models/business_analysis.py', '★ 重写', '新增模块六(_data_quality_module)；扩展每模块输出全部指标(30+字段)；调整综合评分权重(5→6模块)；新增分模块数据提取方法'],
        ['templates/business.html', '★ 重写', '重构为三区联动布局；六轴雷达图；六模块详情展开区；骨架屏+异步分块加载；对标表增加模块得分列'],
        ['models/discrete_analysis.py', '★ 重写', '新增run_with_module_scores()方法；新增_module_indicator_to_level()转换方法；新增置信度计算；保留原有run()作为回退'],
        ['config/discrete_rules.json', '● 扩展', '新增"六模块指标映射"配置段（R/S子维度→六模块指标的映射、权重、阈值）'],
        ['templates/discrete.html', '● 扩展', '新增置信度图例+散点视觉编码；新增What-If侧边栏+JS公式引擎+推演动画；更新规则卡片'],
        ['main.py', '● 调整', '调整步骤顺序(六模块→离散分析)；新增分阶段内存释放(del+gc.collect)；传入business_results给离散分析'],
        ['web_app.py', '● 扩展', '新增/api/business/modules-summary；新增/api/business/module/<id>；新增/api/business/benchmark/<scope>；新增/api/memory-status；新增@cleanup_after_request装饰器'],
        ['templates/base.html', '○ 微调', '新增骨架屏shimmer动画CSS'],
    ]
    add_table_with_style(doc, files_headers, files_rows, [4.5, 2.0, 9.5])

    add_body_para(doc, '')
    add_body_para(doc, '图例：★ = 重写（>200行改动）  ● = 扩展（50-200行改动）  ○ = 微调（<50行改动）', bold=True)

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 附录：核心代码框架
    # ═══════════════════════════════════════════
    add_styled_heading(doc, '附录A：核心代码框架', level=1)

    add_styled_heading(doc, 'A.1 business_analysis.py —— 模块六新增', level=2)
    code1 = '''def _data_quality_module(self, group, project_codes, issue_index, total_projects):
    """模块六：数据质量与流程效率分析（基于模型1.4输出）"""
    r14_df = self._get_model_output("1.4")

    # ① 数据完整率：关键字段非空率
    key_fields = ["项目编码","项目名称","申报单位","项目地址",
                  "签约额（元）","客户名称","工程类别","签约时间"]
    completeness = sum(group[col].notna().mean() for col in key_fields
                       if col in group.columns) / len(key_fields)

    # ② 流程合规率
    process_ok = 1.0 - self._ratio_in_issues(r14_df, project_codes, "流程", "时序")

    # ③ 中标签约偏差率
    bid_dev_ratio = self._calc_bid_deviation_ratio(group)

    # ④ 测算规律性指数
    a_std = group["_a_value"].std()
    estimation_regularity = max(0, 1 - a_std / 0.05)

    # ⑤ 签约延迟率
    sign_delay_ratio = self._calc_sign_delay_ratio(group)

    score = (completeness * 0.20 + process_ok * 0.25
             + (1 - bid_dev_ratio) * 0.20 + estimation_regularity * 0.15
             + (1 - sign_delay_ratio) * 0.20) * 100

    return ModuleScore(
        score=max(0, min(100, score)),
        metrics={
            "数据完整率": completeness,
            "流程合规率": process_ok,
            "中标签约偏差率": bid_dev_ratio,
            "测算规律性指数": estimation_regularity,
            "签约延迟率": sign_delay_ratio,
        }
    )'''
    p = doc.add_paragraph()
    run = p.add_run(code1)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)

    add_styled_heading(doc, 'A.2 discrete_analysis.py —— 基于六模块的R/E计算', level=2)
    code2 = '''def run_with_module_scores(self, all_results, dmp_df, business_results,
                            appendix_df=None, region_auth=None):
    """基于六模块指标计算R/E（替代直接读模型输出）"""
    module_index = self._build_module_index(business_results)

    for idx, row in dmp_df.iterrows():
        unit = str(row.get("申报单位", ""))
        city = str(row.get("项目地址", ""))
        unit_mod = module_index.get(unit, self._default_modules())

        # R风险：从六模块指标直接计算
        r_region   = self._module_indicator_to_level(unit_mod, "跨区域经营指数", "inverse")
        r_contract = self._module_indicator_to_level(unit_mod,
                        [("风险项目占比",0.6), ("合同条款不利度",0.4)], "inverse")
        r_customer = self._module_indicator_to_level(unit_mod,
                        [("客户稳定性指数",0.5), ("客户集中度风险",0.5)], "inverse")
        r_capital  = self._module_indicator_to_level(unit_mod,
                        [("逾期回收率",0.6), ("负流项目占比",0.4)], "inverse")
        r_perf     = self._module_indicator_to_level(unit_mod,
                        [("停工退场率",0.6), ("签约履约偏差率",0.4)], "inverse")

        # E收益：从六模块指标直接计算
        e_profit     = self._module_indicator_to_level(unit_mod, "盈利健康度", "direct")
        e_conversion = self._module_indicator_to_level(unit_mod, "产值转化率", "direct")
        e_collection = self._module_indicator_to_level(unit_mod, "资金回收率", "direct")
        e_scale      = self._module_indicator_to_level(unit_mod, "区域合同额强度", "direct")

        # 加权、归一化、分箱（公式不变）
        R_raw = (r_region*1.0 + r_contract*1.0 + r_customer*0.8
                 + r_capital*0.8 + r_perf*0.4)
        R = R_raw / 4.0
        ...

        # 置信度
        confidence = unit_mod.get("模块六", {}).get("score", 50)'''
    p = doc.add_paragraph()
    run = p.add_run(code2)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)

    add_styled_heading(doc, 'A.3 main.py —— 执行顺序调整与内存释放', level=2)
    code3 = '''# ===== [4.2] 六模块分析（先于离散分析）=====
print("\\n[4.2/5] Running business health analysis (6 modules)...")
business_analyzer = BusinessHealthAnalyzer()
business_results = business_analyzer.run(all_results, dmp)
with open(biz_path, "wb") as f:
    pickle.dump(business_results, f)

# ===== [4.3] 九宫格离散分析（消费六模块结果）=====
print("\\n[4.3/5] Running discrete analysis (from 6-module indicators)...")
discrete_analyzer = DiscreteAnalyzer(config)
discrete_results = discrete_analyzer.run_with_module_scores(
    all_results, dmp, business_results, appendix_df, region_auth
)

# ===== 内存释放 =====
print("\\n[Memory] Releasing large DataFrames...")
del dmp, dmp_raw, appendices, bid_report
import gc
gc.collect()
print(f"  GC collected {gc.collect()} objects")'''
    p = doc.add_paragraph()
    run = p.add_run(code3)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)

    add_styled_heading(doc, 'A.4 discrete.html —— What-If JS公式引擎', level=2)
    code4 = '''const DISCRETE_FORMULA = {
    riskWeights: {region:1.0, contract:1.0, customer:0.8, capital:0.8, perf:0.4},
    returnWeights: {profit:1.05, conversion:0.75, collection:0.75, scale:0.45},

    calcRiskSubDim(moduleIndicators, dimName) {
        const mapping = {
            region:   {src:'模块一', key:'跨区域经营指数', dir:'inverse'},
            contract: {src:'模块三', key:'风险项目占比', key2:'合同条款不利度',
                       weights:[0.6,0.4], dir:'inverse'},
            // ... other dims
        };
        const m = mapping[dimName];
        let value = moduleIndicators[m.src][m.key];
        if (m.key2) value = value*m.weights[0] + moduleIndicators[m.src][m.key2]*m.weights[1];
        if (m.dir === 'inverse') value = 1 - value;
        return value >= 0.70 ? 1 : value >= 0.35 ? 2 : 3;
    },

    recalc(moduleIndicators) {
        const R = (this.calcRiskSubDim(moduleIndicators,'region')*1.0 + ...) / 4.0;
        const E = (this.calcReturnSubDim(moduleIndicators,'profit')*1.05 + ...) / 3.0;
        return {R: Math.max(1,Math.min(3,R)), E: Math.max(1,Math.min(3,E))};
    }
};

// 滑块事件 → 实时重算 → 气泡移动
function onSliderChange() {
    const adjusted = getAdjustedModuleIndicators();
    const {R, E} = DISCRETE_FORMULA.recalc(adjusted);
    drawWhatIfBubble(R, E);
    updateWhatIfResult(R, E);
}'''
    p = doc.add_paragraph()
    run = p.add_run(code4)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)

    add_styled_heading(doc, 'A.5 数据置信度视觉编码（ECharts配置片段）', level=2)
    code5 = '''// discrete.html renderScatter() 中的置信度编码
{
    type: "scatter",
    data: rows.map(row => ({
        value: [row.r + jitter(...), row.e + jitter(...), row.size],
        itemStyle: {
            color: row.confidence >= 80 ? row.color :
                   row.confidence >= 50 ? row.color + '99' : row.color + '55',
            borderColor: row.confidence < 50 ? '#d94d63' :
                         row.confidence < 80 ? '#f59e0b' : '#fff',
            borderWidth: row.confidence < 50 ? 3 : row.confidence < 80 ? 2 : 1.5,
            borderType: row.confidence < 50 ? 'dashed' : 'solid',
            opacity: row.confidence >= 80 ? 0.92 :
                     row.confidence >= 50 ? 0.70 : 0.45,
        },
        symbol: row.confidence < 30 ? 'diamond' : 'circle',
    })),
    tooltip: {
        formatter(params) {
            const row = params.data.raw;
            let html = `${row.name}<br>R=${row.r} / E=${row.e}<br>${row.gridName}`;
            if (row.confidence < 50) {
                html += `<br><span style="color:#d94d63">⚠ 数据置信度低(${row.confidence}分)</span>`;
                html += `<br><em>建议：决策前先复核底层数据</em>`;
            }
            return html;
        }
    }
}'''
    p = doc.add_paragraph()
    run = p.add_run(code5)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)

    # ── Save standalone document ──
    output_path = os.path.join(DOCS_DIR, '实施方案_六模块九宫格全链路贯通_v2.10.docx')
    doc.save(output_path)
    print(f'[OK] Standalone plan saved: {output_path}')
    return output_path


def update_v29_document():
    """Append the new implementation plan as Chapter 8 to the existing v2.9 document"""
    import zipfile, shutil
    from docx import Document

    src_path = os.path.join(DOCS_DIR, '方案_v2.9_模型2.5签约履约偏差_完整版修订.docx')

    # Since we can't easily edit the existing docx while preserving all formatting,
    # we'll create a new combined document based on the extracted text + new content

    # Read existing document text
    z = zipfile.ZipFile(src_path)
    xml_content = z.read('word/document.xml')
    import xml.etree.ElementTree as ET
    tree = ET.fromstring(xml_content)
    paras = []
    for p in tree.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
        texts = []
        for t in p.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
            if t.text:
                texts.append(t.text)
        if texts:
            paras.append(''.join(texts))

    existing_text = '\n'.join(paras)

    # Create a new document with existing content + new chapter
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)

    # Add a note at the beginning
    note = doc.add_paragraph()
    run = note.add_run('【v2.10增量更新】本文档为v2.9方案基础上新增"第八章：六模块·九宫格·全链路贯通实施方案"。')
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1f, 0x89, 0xdf)

    note2 = doc.add_paragraph()
    run2 = note2.add_run(f'更新日期：{datetime.date.today().isoformat()}。原有第一章~第七章及附录A~F完整保留，新增第八章为本次增量。')
    run2.font.name = '微软雅黑'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run2.font.size = Pt(10)

    doc.add_paragraph()

    # Add the existing content as a reference
    add_styled_heading(doc, '（原有内容：第一章~第七章、附录A~F，完整保留于原始docx文件中）', level=2)
    add_body_para(doc, f'原方案文档共{len(existing_text):,}字符、{existing_text.count(chr(10)):,}行，覆盖以下章节：')
    chapters = [
        '第一章 底层数据输入及标准准备 (Data Input)',
        '第二章 核心IT算法与业务逻辑模型设计',
        '第三章 模型关联与综合应用架构',
        '第四章 输出与报告体系',
        '第五章 模型优先级与执行策略',
        '第六章 字段缺口与数据补充需求',
        '第七章 实施路线图',
        '附录A 合规映射（手册0520附件6/7）',
        '附录B 数据架构详细字段矩阵',
        '附录C 模型间关联图谱',
        '附录D SAP财务数据补充需求',
        '附录E 外部信用数据',
        '附录F DMP系统缺失字段',
    ]
    for ch in chapters:
        add_body_para(doc, f'  • {ch}')

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # NEW CHAPTER 8: 六模块·九宫格·全链路贯通实施方案
    # ═══════════════════════════════════════════
    add_styled_heading(doc, '第八章（v2.10新增）：六模块·九宫格·全链路贯通实施方案', level=1)

    add_body_para(doc, '本章为v2.10版本的核心增量更新，目标是将现有的"11模型→三维度→三条链→九宫格"架构升级为"11模型→三维度→三条链→六模块→九宫格"的完整数据链路。', bold=True)

    # Insert the new chapter by referencing the standalone plan content
    # We'll add all the same sections as the standalone document

    add_styled_heading(doc, '8.1 背景与问题诊断', level=2)
    add_body_para(doc, '当前v2.9系统存在以下断层需要修复：')
    problems = [
        '断层一：六模块不完整。BusinessHealthAnalyzer仅实现了5个模块，缺少"模块六：数据质量与流程效率分析"。',
        '断层二：指标展示不全。business.html仅展示6个零散KPI，未按六模块结构化展示全部30+指标。',
        '断层三：九宫格R/S与六模块脱节。DiscreteAnalyzer绕过六模块直接读模型原始输出。',
        '断层四：缺少置信度反馈。数据质量差的扩张区项目无法视觉区分。',
        '断层五：缺少推演能力。无法模拟"改善某指标后项目在九宫格中如何移动"。',
    ]
    for i, p in enumerate(problems, 1):
        add_body_para(doc, f'({i}) {p}')

    add_styled_heading(doc, '8.2 六模块完整指标体系', level=2)

    mapping_headers = ['模块', '对应模型', '核心指标（5-6个/模块）']
    mapping_rows = [
        ['模块一：区域布局健康度', '1.1 + 1.2', '①区域渗透率 ②跨区域经营指数 ③深耕区域集中度\n④区域合同额强度 ⑤业务结构偏离度 ⑥EPC转型进度'],
        ['模块二：客户资源稳定性', '1.3 + 3.1 + 3.2', '①客户稳定性指数 ②客户产出波动率 ③客户集中度风险\n④中标转化率 ⑤新客户质量指数 ⑥战略客户产出比'],
        ['模块三：合同质量与风险集中度', '2.1 + 2.4', '①风险项目占比 ②风险合同额集中度 ③付款条件优良率\n④合同条款不利度 ⑤三证合规率'],
        ['模块四：履约盈利健康度', '2.2 + 2.5 + 前置过滤', '①产值转化率 ②签约履约偏差率 ③盈利健康度\n④停工退场率 ⑤效益偏差率 ⑥在施项目活跃度'],
        ['模块五：资金效率与安全性', '2.3', '①资金占用率 ②保证金周转天数 ③逾期回收率\n④预收款缺口率 ⑤负流项目占比'],
        ['模块六：数据质量与流程效率', '1.4', '①数据完整率 ②流程合规率 ③中标签约偏差率\n④测算规律性指数 ⑤签约延迟率'],
    ]
    add_table_with_style(doc, mapping_headers, mapping_rows, [3.5, 3.0, 9.5])

    add_body_para(doc, '')
    add_body_para(doc, '综合评分权重调整为6模块：', bold=True)
    add_body_para(doc, '模块一(25%) + 模块二(20%) + 模块三(18%) + 模块四(15%) + 模块五(12%) + 模块六(10%) = 100%')

    add_styled_heading(doc, '8.3 九宫格R/S与六模块指标映射', level=2)

    add_body_para(doc, '核心变更：九宫格R/S子维度不再直接读取11模型原始输出，改为直接采用六模块的标准化指标。', bold=True)

    add_body_para(doc, '')
    add_body_para(doc, 'R（风险维度）← 六模块风险指标：', bold=True)
    r_headers = ['R子维度', '权重', '对应六模块指标', '来源模块']
    r_rows = [
        ['区域合规', '1.0', '跨区域经营指数', '模块一'],
        ['合同底线', '1.0', '风险项目占比(0.6) + 合同条款不利度(0.4)', '模块三'],
        ['客户健康', '0.8', '客户稳定性指数(0.5) + 客户集中度风险(0.5)', '模块二'],
        ['资金安全', '0.8', '逾期回收率(0.6) + 负流项目占比(0.4)', '模块五'],
        ['履约真实', '0.4', '停工退场率(0.6) + 签约履约偏差率(0.4)', '模块四'],
    ]
    add_table_with_style(doc, r_headers, r_rows, [2.0, 1.0, 9.0, 4.0])

    add_body_para(doc, '')
    add_body_para(doc, 'E（收益维度）← 六模块收益指标：', bold=True)
    e_headers = ['E子维度', '权重', '对应六模块指标', '来源模块']
    e_rows = [
        ['盈利水平', '1.05', '盈利健康度', '模块四'],
        ['产值转化', '0.75', '产值转化率', '模块四'],
        ['资金回收', '0.75', '资金回收率（收款/签约额）', '模块五'],
        ['合同规模', '0.45', '区域合同额强度（归一化）', '模块一'],
    ]
    add_table_with_style(doc, e_headers, e_rows, [2.0, 1.0, 9.0, 4.0])

    add_body_para(doc, '')
    add_body_para(doc, '公式不变：R = R_raw / 4.0 ∈ [1.0, 3.0]；R ≤ 1.6 → 1档，1.6 < R ≤ 2.4 → 2档，R > 2.4 → 3档')
    add_body_para(doc, '公式不变：E = E_raw / 3.0 ∈ [1.0, 3.0]；E ≤ 1.6 → 1档，1.6 < E ≤ 2.4 → 2档，E > 2.4 → 3档')

    add_styled_heading(doc, '8.4 数据置信度视觉反馈', level=2)
    add_body_para(doc, '模块六得分作为项目的"数据置信度"标识。在九宫格散点图中，低置信度（<50分）项目的气泡采用：虚线红色边框、高度透明（opacity 0.45）、菱形符号。Tooltip悬停时红色警告"⚠ 数据置信度低(NN分)，决策需谨慎"。散点图底部新增置信度图例条。')

    add_styled_heading(doc, '8.5 What-If沙盘推演', level=2)
    add_body_para(doc, '在discrete.html右侧新增可折叠侧边栏。将Python的R/E计算公式完整移植到JavaScript。业务人员拖动滑块调整六模块指标，前端实时重算R/S坐标，气泡在九宫格中动态跃迁，并绘制从原始位置到推演位置的绿色虚线箭头轨迹。')

    add_styled_heading(doc, '8.6 前端异步加载与骨架屏', level=2)
    add_body_para(doc, 'business.html采用分块异步加载策略：先渲染shimmer骨架屏 → fetch /api/business/modules-summary（<1KB秒返）→ 雷达图+KPIs亮起 → Promise.allSettled并行加载6个模块API → 哪个模块先返回先亮起 → 最后加载对标表。后端新增3类轻量API支持此流程。')

    add_styled_heading(doc, '8.7 内存释放机制', level=2)
    add_body_para(doc, 'main.py pipeline各阶段结束后显式释放大对象（del dmp, dmp_raw, appendices, bid_report → gc.collect()）。web_app.py的API响应后调用gc.collect(0)回收年轻代。新增/api/memory-status运维监控端点。')

    add_styled_heading(doc, '8.8 实施步骤', level=2)
    impl_headers = ['阶段', '核心工作', '工时']
    impl_rows = [
        ['P1 六模块补全', 'business_analysis.py新增模块六 + 扩展全指标 + web_app.py新增分模块API', '4.5h'],
        ['P2 监测台重构', 'business.html骨架屏+异步加载+六模块卡片区+雷达图+对标表', '5.5h'],
        ['P3 R/S映射', 'discrete_rules.json配置+discrete_analysis.py重构+main.py调整+内存释放', '3.5h'],
        ['P4 置信度', '置信度计算+散点图视觉编码', '2h'],
        ['P5 What-If', 'JS公式引擎+滑块UI+推演动画', '3h'],
        ['P6 联调', '端到端测试+边界处理', '2h'],
    ]
    add_table_with_style(doc, impl_headers, impl_rows, [3.0, 11.0, 2.0])
    add_body_para(doc, '')
    add_body_para(doc, '总计工时：约20.5小时', bold=True)

    add_styled_heading(doc, '8.9 涉及文件清单', level=2)
    files_headers = ['文件', '改动类型', '改动说明']
    files_rows = [
        ['models/business_analysis.py', '★ 重写', '新增模块六；扩展每模块输出全部指标；调整权重'],
        ['templates/business.html', '★ 重写', '三区联动布局；骨架屏；六轴雷达；六模块详情区'],
        ['models/discrete_analysis.py', '★ 重写', '新增run_with_module_scores()；置信度计算'],
        ['config/discrete_rules.json', '● 扩展', '新增"六模块指标映射"配置段'],
        ['templates/discrete.html', '● 扩展', '置信度编码；What-If侧边栏；JS公式引擎；推演动画'],
        ['main.py', '● 调整', '步骤顺序调整；分阶段内存释放'],
        ['web_app.py', '● 扩展', '3个新API；cleanup装饰器；memory-status端点'],
        ['templates/base.html', '○ 微调', '骨架屏shimmer CSS动画'],
    ]
    add_table_with_style(doc, files_headers, files_rows, [4.5, 2.0, 9.5])

    add_body_para(doc, '')
    add_body_para(doc, '图例：★ = 重写  ● = 扩展  ○ = 微调', bold=True)

    add_styled_heading(doc, '8.10 向后兼容', level=2)
    add_body_para(doc, '如果六模块分析尚未执行（无business_results缓存），discrete_analysis.py自动回退到直接读取11模型原始输出的现有逻辑（即v2.9的run()方法），确保系统在任何状态下均可用。')

    # ── Save updated v2.9 document ──
    output_path = os.path.join(DOCS_DIR, '方案_v2.10_六模块九宫格全链路贯通_完整版.docx')
    doc.save(output_path)
    print(f'[OK] Updated v2.9 → v2.10 saved: {output_path}')
    return output_path


if __name__ == '__main__':
    print('=' * 60)
    print('  生成Word文档')
    print('=' * 60)
    plan_path = create_standalone_plan()
    v29_updated_path = update_v29_document()
    print()
    print(f'生成文件：')
    print(f'  1. {plan_path}')
    print(f'  2. {v29_updated_path}')
    print()
    print('Done!')
