import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

with open(r"c:\Users\sasa\Desktop\模型建设\模型2：市场营销\附件5.md", "r", encoding="utf-8") as f:
    content = f.read()

doc = Document()
st = doc.styles["Normal"]
st.font.name = "宋体"; st.font.size = Pt(12)
st.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

def T(t):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t); r.font.name = "黑体"; r.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(22); r.bold = True

def H(t):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.font.name = "宋体"; r.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(16); r.bold = True; p.paragraph_format.space_before = Pt(18)

def S(t):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.font.name = "宋体"; r.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(14); r.bold = True; p.paragraph_format.space_before = Pt(12)

def B(segs, sz=12):
    p = doc.add_paragraph()
    for x, b in segs:
        r = p.add_run(x); r.font.name = "宋体"; r.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(sz); r.bold = b
    p.paragraph_format.line_spacing = 1.5

def Q(t):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.font.name = "宋体"; r.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(10); r.italic = True; r.font.color.rgb = RGBColor(0x66, 0x7B, 0x90)

def TB(hd, rows):
    t = doc.add_table(rows=min(len(rows)+1, 35), cols=len(hd)); t.style = "Table Grid"
    for i, h in enumerate(hd):
        c = t.rows[0].cells[i]; c.text = ""
        r = c.paragraphs[0].add_run(h); r.font.name = "宋体"; r.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(9); r.bold = True
    for ri, row in enumerate(rows[:34]):
        for ci, v in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ""
            r = c.paragraphs[0].add_run(str(v)[:140]); r.font.name = "宋体"; r.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            r.font.size = Pt(9)
    doc.add_paragraph()

def P(s):
    pts = re.split(r"(\*\*.*?\*\*)", s)
    return [(p[2:-2], True) if p.startswith("**") and p.endswith("**") else (p, False) for p in pts]

lns = content.split("\n")
i, it, tr, th, iq, ql = 0, False, [], [], False, []
while i < len(lns):
    s = lns[i].strip()
    if not s:
        if iq and ql: Q(" ".join(ql)); ql = []; iq = False
        i += 1; continue
    if s.startswith("> "): ql.append(s[2:]); iq = True; i += 1; continue
    if s.startswith("```") or s == "---": i += 1; continue
    if "|" in s and s.count("|") >= 2:
        if not it:
            it = True; th = [h.strip() for h in s.split("|")[1:-1]]; tr = []
        elif not all(c in "- :" for c in s.replace("|", "").strip()):
            cl = [h.strip() for h in s.split("|")[1:-1]]
            if cl: tr.append(cl)
        i += 1; continue
    else:
        if it and th: TB(th, tr); th, tr, it = [], [], False
    if s.startswith("# ") and not s.startswith("## "): T(s[2:])
    elif s.startswith("## "): H(s[3:])
    elif s.startswith("### "): S(s[4:])
    elif s.startswith("- ") or re.match(r"^\d+\.\s", s): B(P(re.sub(r"^\d+\.\s|^-\s", "", s)))
    else: B(P(s))
    i += 1
if it and th: TB(th, tr)
for sec in doc.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

out = r"c:\Users\sasa\Desktop\模型建设\模型2：市场营销\附件5_更新版_v9.docx"
doc.save(out)
print(f"v9 saved: {out}")
