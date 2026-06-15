"""
Generate two Word documents (v2 — with 4 supplementary sections):
1. 实施方案_六模块九宫格全链路贯通_v2.10.docx — standalone
2. 方案_v2.10_六模块九宫格全链路贯通_完整版.docx — v2.9 upgrade
"""
import os, sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

PROJECT_ROOT = r"C:\Users\sasa\Desktop\模型建设\模型2：市场营销"
DOCS_DIR = os.path.join(PROJECT_ROOT, "_documents")

# ── helpers ──────────────────────────────────────────────
def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def add_para(doc, text, bold=False, size=11, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def add_code(doc, code, size=8):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(size)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = '微软雅黑'; run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(10); run.bold = True
        run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1f89df')
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = '微软雅黑'; run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.size = Pt(9)
            if r % 2 == 1: set_cell_shading(cell, 'F5F8FC')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = '微软雅黑'; run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    return p

def add_alert(doc, text, level='danger'):
    """Highlighted alert box (simulated with colored paragraph)"""
    colors = {'danger': RGBColor(0xd9,0x4d,0x63), 'warning': RGBColor(0xb7,0x7b,0x12),
              'info': RGBColor(0x1f,0x89,0xdf)}
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'; run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(10); run.bold = True
    run.font.color.rgb = colors.get(level, RGBColor(0,0,0))
    return p


# ═══════════════════════════════════════════════════════════
# DOCUMENT 1: STANDALONE IMPLEMENTATION PLAN
# ═══════════════════════════════════════════════════════════
def create_standalone_plan():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

    # Title
    t = doc.add_heading('全面数字化营销审计系统', level=0)
    for r in t.runs: r.font.name = '微软雅黑'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    st = doc.add_heading('六模块·九宫格·全链路贯通实施方案', level=1)
    for r in st.runs: r.font.name = '微软雅黑'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    meta = doc.add_paragraph()
    for text, bold in [
        ('版本：v2.10（基于v2.9增量升级）\n', True),
        (f'日期：{datetime.date.today().isoformat()}\n', False),
        ('状态：待实施  |  关联：方案_v2.9_模型2.5签约履约偏差_完整版修订.docx', False),
    ]:
        run = meta.add_run(text); run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑'); run.font.size = Pt(10); run.bold = bold

    doc.add_page_break()

    # ═══════════════ CHAPTER 1 ═══════════════
    add_heading(doc, '第一章 背景与目标', 1)
    add_heading(doc, '1.1 问题诊断', 2)
    add_para(doc, '当前v2.9系统存在以下断层：')
    for i, p in enumerate([
        '断层一：六模块不完整。BusinessHealthAnalyzer仅实现5个模块，缺"模块六：数据质量与流程效率分析"。',
        '断层二：指标展示不全。business.html仅展示6个零散KPI，未按六模块结构化展示30+指标。',
        '断层三：九宫格R/S与六模块脱节。DiscreteAnalyzer直接读模型原始输出，绕过六模块指标层。',
        '断层四：缺少置信度反馈。即使项目落入"扩张区"，若底层数据质量差，决策者无从知晓。',
        '断层五：缺少推演能力。业务人员无法模拟"如果改善某个指标，项目会在九宫格中如何移动"。',
        '断层六：前端规则引擎硬编码。What-If公式若硬编码JS，与后端discrete_rules.json不同步时将产生撕裂。',
    ], 1):
        add_para(doc, f'({i}) {p}')

    add_heading(doc, '1.2 目标架构', 2)
    add_para(doc, '建立"11模型→三维度→三条链→六模块→九宫格"的完整链路，六模块为数据中枢，所有上层分析均从六模块取数。')
    add_para(doc, '四项关键架构约束：')
    add_bullet(doc, '规则引擎单源：所有公式权重由 /api/discrete/config 统一下发，前后端同源。')
    add_bullet(doc, '数据容错防线：Module 4聚合时保留真实负向差异，abs()剥离成果无损穿透。')
    add_bullet(doc, '导出对齐锁：Excel生成器绑定动态行号索引，防止GC并发下窜行。')
    add_bullet(doc, '视觉质感统一：全站毛玻璃设计系统，减少复杂数据的视觉负担。')

    doc.add_page_break()

    # ═══════════════ CHAPTER 2: 六模块指标体系 ═══════════════
    add_heading(doc, '第二章 六模块完整指标体系', 1)
    add_heading(doc, '2.1 11模型→六模块映射', 2)
    add_table(doc,
        ['模块', '对应模型', '核心指标（5-6个/模块）'],
        [['模块一：区域布局健康度','1.1+1.2','①区域渗透率 ②跨区域经营指数 ③深耕区域集中度\n④区域合同额强度 ⑤业务结构偏离度 ⑥EPC转型进度'],
         ['模块二：客户资源稳定性','1.3+3.1+3.2','①客户稳定性指数 ②客户产出波动率 ③客户集中度风险\n④中标转化率 ⑤新客户质量指数 ⑥战略客户产出比'],
         ['模块三：合同质量与风险集中度','2.1+2.4','①风险项目占比 ②风险合同额集中度 ③付款条件优良率\n④合同条款不利度 ⑤三证合规率'],
         ['模块四：履约盈利健康度','2.2+2.5+前置过滤','①产值转化率 ②签约履约偏差率 ③盈利健康度\n④停工退场率 ⑤效益偏差率 ⑥在施项目活跃度'],
         ['模块五：资金效率与安全性','2.3','①资金占用率 ②保证金周转天数 ③逾期回收率\n④预收款缺口率 ⑤负流项目占比'],
         ['模块六：数据质量与流程效率','1.4','①数据完整率 ②流程合规率 ③中标签约偏差率\n④测算规律性指数 ⑤签约延迟率']],
        [3.5,3.0,9.5])

    add_heading(doc, '2.2 每个模块指标详细定义', 2)

    # Module 1
    add_heading(doc, '模块一：区域布局健康度（模型1.1+1.2）', 3)
    add_table(doc, ['指标名称','计算逻辑','经营意义'],
        [['区域渗透率','实际签约城市数/授权城市数×100%','衡量区域深耕深度'],
         ['跨区域经营指数','非常规区域合同额/总合同额','反映经营扩张激进程度'],
         ['深耕区域集中度','核心+重点城市合同额占比','评估区域聚焦度'],
         ['区域合同额强度','单城市平均合同额','识别"撒胡椒面"式布局'],
         ['业务结构偏离度','Σ|实际占比-战略占比|/板块数','衡量战略执行偏差'],
         ['EPC转型进度','EPC模式合同额/总合同额','业务模式升级速度']], [3.5,6.0,6.5])

    # Module 2
    add_heading(doc, '模块二：客户资源稳定性（模型1.3+3.1+3.2）', 3)
    add_table(doc, ['指标名称','计算逻辑','经营意义'],
        [['客户稳定性指数','1-(流失客户数/上年客户总数)','客户资源流失风险'],
         ['客户产出波动率','标准差(年度合同额)/均值','客户产出是否稳定'],
         ['客户集中度风险','前5大客户合同额占比','依赖度过高预警'],
         ['中标转化率','签约额/中标额','营销漏斗效率'],
         ['新客户质量指数','国企/政府类新客户占比×粘性系数','获客质量评估'],
         ['战略客户产出比','战略客户合同额/总合同额','核心客户贡献度']], [3.5,6.0,6.5])

    # Module 3
    add_heading(doc, '模块三：合同质量与风险集中度（模型2.1+2.4）', 3)
    add_table(doc, ['指标名称','计算逻辑','经营意义'],
        [['风险项目占比','触碰严禁投标底线项目数/总项目数','整体风险水位'],
         ['风险合同额集中度','高风险项目合同额/总合同额','风险敞口规模'],
         ['付款条件优良率','付款达标项目数/总项目数','资金回收保障度'],
         ['合同条款不利度','高风险条款项目数/总项目数','合同谈判质量'],
         ['三证合规率','三证齐全项目数/已开工项目数','合规经营水平']], [3.5,6.0,6.5])

    # Module 4 (with data defense note)
    add_heading(doc, '模块四：履约盈利健康度（模型2.2+2.5+前置过滤）', 3)
    add_table(doc, ['指标名称','计算逻辑','经营意义'],
        [['产值转化率','实际完成产值/签约额','合同落地效率'],
         ['签约履约偏差率','1-产值转化率(签约>12月)','识别"签而不做"项目'],
         ['盈利健康度','A值≥底线项目占比','整体盈利质量'],
         ['停工退场率','停工退场项目数/已开工项目数','项目夭折风险'],
         ['效益偏差率','|A值-实际利润率|>1%项目占比','测算准确性'],
         ['在施项目活跃度','在施项目产值>0且收款>0占比','在建项目健康度']], [3.5,6.0,6.5])

    add_alert(doc, '⚠ 数据容错防线（详见第十一章）：模块四上卷聚合时，DMP的材料抵消与退场核销逻辑会产生负向产值数据。必须严格保留真实负向差异，确保abs()剥离成果无损穿透至盈利健康度计算，防止成本双倍扣减导致模块得分断崖式失真。', 'danger')

    # Module 5
    add_heading(doc, '模块五：资金效率与安全性（模型2.3）', 3)
    add_table(doc, ['指标名称','计算逻辑','经营意义'],
        [['资金占用率','(保证金+预收款应收)/总合同额','资金沉淀规模'],
         ['保证金周转天数','平均(实际回收日-约定退还日)','资金回收效率'],
         ['逾期回收率','逾期>90天款项金额/应回收总金额','资金风险敞口'],
         ['预收款缺口率','(应收预收款-实收预收款)/应收预收款','预收款回收缺口'],
         ['负流项目占比','资金结余<<0项目数/总项目数','现金流健康度']], [3.5,6.0,6.5])

    # Module 6
    add_heading(doc, '模块六：数据质量与流程效率（模型1.4）', 3)
    add_table(doc, ['指标名称','计算逻辑','经营意义'],
        [['数据完整率','关键字段非空率(DMP 78字段)','数据治理水平'],
         ['流程合规率','招文评审→交标→中标→签约时序合规占比','流程执行质量'],
         ['中标签约偏差率','|中标额-签约额|>5%项目占比','报价稳定性'],
         ['测算规律性指数','同单位利润率标准差','数据真实性'],
         ['签约延迟率','预计签约日已过未签约项目占比','营销转化效率']], [3.5,6.0,6.5])

    add_heading(doc, '2.3 综合评分权重', 2)
    add_table(doc, ['模块','原权重','新权重','变化'],
        [['模块一：区域布局','30%','25%','分走5%给模块六'],
         ['模块二：客户稳定','20%','20%','不变'],
         ['模块三：合同质量','20%','18%','分走2%给模块六'],
         ['模块四：履约盈利','15%','15%','不变'],
         ['模块五：资金效率','15%','12%','分走3%给模块六'],
         ['模块六：数据质量','—','10%','新增']], [4.0,2.5,2.5,7.0])
    add_para(doc, '综合得分 = 模块一(25%) + 模块二(20%) + 模块三(18%) + 模块四(15%) + 模块五(12%) + 模块六(10%)', bold=True)
    add_para(doc, '分区阈值：≥80强势区 | 65-79稳健区 | <65承压区')

    doc.add_page_break()

    # ═══════════════ CHAPTER 3: R/S映射 ═══════════════
    add_heading(doc, '第三章 九宫格R/S与六模块指标映射', 1)
    add_heading(doc, '3.1 核心原则', 2)
    add_para(doc, '九宫格R/S所有子维度不再直接读取11模型原始输出，改为直接采用六模块标准化指标。监测台与九宫格指标完全同源。')

    add_heading(doc, '3.2 R（风险维度）← 六模块指标', 2)
    add_table(doc, ['R子维度','权重','六模块指标','模块','方向','分箱阈值'],
        [['区域合规','1.0','跨区域经营指数','模块一','inverse','≤20%→1,20-50%→2,>50%→3'],
         ['合同底线','1.0','风险项目占比(0.6)+合同条款不利度(0.4)','模块三','inverse','≤15%→1,15-35%→2,>35%→3'],
         ['客户健康','0.8','客户稳定性指数(0.5)+客户集中度风险(0.5)','模块二','inverse','稳定性<60%→3,集中度>60%→3'],
         ['资金安全','0.8','逾期回收率(0.6)+负流项目占比(0.4)','模块五','inverse','逾期>20%→3,负流>15%→3'],
         ['履约真实','0.4','停工退场率(0.6)+签约履约偏差率(0.4)','模块四','inverse','停工>10%→3,偏差>50%→3']],
        [2.0,1.0,5.0,1.5,1.5,5.0])
    add_para(doc, 'R_raw = 区域合规×1.0 + 合同底线×1.0 + 客户健康×0.8 + 资金安全×0.8 + 履约真实×0.4')
    add_para(doc, 'R = R_raw / 4.0 ∈ [1.0, 3.0]；R≤1.6→低风险(1)；1.6<R≤2.4→中风险(2)；R>2.4→高风险(3)')

    add_heading(doc, '3.3 E（收益维度）← 六模块指标', 2)
    add_table(doc, ['E子维度','权重','六模块指标','模块','方向','分箱阈值'],
        [['盈利水平','1.05','盈利健康度','模块四','direct','≥80%→3,50-80%→2,<50%→1'],
         ['产值转化','0.75','产值转化率','模块四','direct','≥80%→3,50-80%→2,<50%→1'],
         ['资金回收','0.75','资金回收率(收款/签约额)','模块五','direct','≥60%→3,30-60%→2,<30%→1'],
         ['合同规模','0.45','区域合同额强度(归一化)','模块一','direct','≥5亿→3,1-5亿→2,<1亿→1']],
        [2.0,1.0,5.0,1.5,1.5,5.0])
    add_para(doc, 'E_raw = 盈利水平×1.05 + 产值转化×0.75 + 资金回收×0.75 + 合同规模×0.45')
    add_para(doc, 'E = E_raw / 3.0 ∈ [1.0, 3.0]；E≤1.6→低收益(1)；1.6<E≤2.4→中收益(2)；E>2.4→高收益(3)')

    add_heading(doc, '3.4 九宫格策略矩阵', 2)
    add_table(doc, ['格子','风险-收益','命名','特征','策略'],
        [['(1,1)','低-低','退出区','深耕区域但项目小、盈利差','主动退出或合并，释放资源'],
         ['(1,2)','低-中','培育区','合规经营但规模未起量','资源倾斜，扩大市场份额'],
         ['(1,3)','低-高','扩张区','深耕+大项目+高盈利','复制推广，设立区域中心'],
         ['(2,1)','中-低','观察区','非常规+小项目+盈利差','收缩投入，限期改善'],
         ['(2,2)','中-中','维持区','单一风险但收益尚可','常规管理，动态监控'],
         ['(2,3)','中-高','优化区','非常规区域但项目优质','压缩风险敞口，推动合规化'],
         ['(3,1)','高-低','淘汰区','违规+红线+亏损','立即止损，启动问责'],
         ['(3,2)','高-中','整顿区','红线但收益尚可','限期整改，回溯审批'],
         ['(3,3)','高-高','警惕区','刀口舔血：违规但高盈利','专人周跟踪，一案一策']],
        [1.2,1.5,2.0,5.3,6.0])

    doc.add_page_break()

    # ═══════════════ CHAPTER 4: 规则引擎双源一致性 ═══════════════
    add_heading(doc, '第四章 规则引擎双源一致性保障（防What-If致命伤）', 1)

    add_alert(doc, '⚠ 架构反模式警告：如果在JS中硬编码DISCRETE_FORMULA权重和映射关系，一旦业务部门调整config/discrete_rules.json而开发者忘记同步修改前端JS，沙盘推演结果将与后端实际运算彻底撕裂。这是What-If模式的致命伤。', 'danger')

    add_heading(doc, '4.1 补强方案：/api/discrete/config 动态规则下发', 2)
    add_para(doc, '绝对禁止在JavaScript中硬编码任何权重、阈值或映射关系。改为由后端API统一下发：')
    add_para(doc, '')
    add_para(doc, '新增API端点：', bold=True)
    add_para(doc, 'GET /api/discrete/config  ——  返回discrete_rules.json解析后的完整规则对象')

    add_para(doc, '')
    add_para(doc, '返回数据结构：', bold=True)
    add_code(doc, '''{
  "riskWeights": {"region": 1.0, "contract": 1.0, "customer": 0.8, "capital": 0.8, "perf": 0.4},
  "returnWeights": {"profit": 1.05, "conversion": 0.75, "collection": 0.75, "scale": 0.45},
  "riskMapping": {
    "region":   {"module": 1, "indicators": [["跨区域经营指数", 1.0]], "direction": "inverse", "thresholds": [0.20, 0.50]},
    "contract": {"module": 3, "indicators": [["风险项目占比", 0.6], ["合同条款不利度", 0.4]], "direction": "inverse", "thresholds": [0.15, 0.35]},
    "customer": {"module": 2, "indicators": [["客户稳定性指数", 0.5], ["客户集中度风险", 0.5]], "direction": "inverse", "thresholds": [0.40, 0.60]},
    "capital":  {"module": 5, "indicators": [["逾期回收率", 0.6], ["负流项目占比", 0.4]], "direction": "inverse", "thresholds": [0.10, 0.25]},
    "perf":     {"module": 4, "indicators": [["停工退场率", 0.6], ["签约履约偏差率", 0.4]], "direction": "inverse", "thresholds": [0.05, 0.20]}
  },
  "returnMapping": {
    "profit":     {"module": 4, "indicators": [["盈利健康度", 1.0]], "direction": "direct", "thresholds": [0.50, 0.80]},
    "conversion": {"module": 4, "indicators": [["产值转化率", 1.0]], "direction": "direct", "thresholds": [0.50, 0.80]},
    "collection": {"module": 5, "indicators": [["资金回收率", 1.0]], "direction": "direct", "thresholds": [0.30, 0.60]},
    "scale":      {"module": 1, "indicators": [["区域合同额强度", 1.0]], "direction": "direct", "thresholds": [0.20, 0.60]}
  },
  "gridStrategies": { ... },
  "cutThresholds": {"riskLow": 1.6, "riskHigh": 2.4, "returnLow": 1.6, "returnHigh": 2.4}
}''')

    add_para(doc, '')
    add_para(doc, '后端实现（web_app.py）：', bold=True)
    add_code(doc, '''@app.route("/api/discrete/config")
def api_discrete_config():
    """返回完整的离散化规则配置，供前端What-If沙盘动态计算使用"""
    disc_path = PROJECT_ROOT / "config" / "discrete_rules.json"
    if not disc_path.exists():
        return jsonify({"error": "discrete_rules.json not found"}), 404
    with open(disc_path, "r", encoding="utf-8") as f:
        raw = json.load(f)
    rules = raw.get("离散化分析", {})

    # 构建前端可直接消费的简化结构
    risk_cfg = rules.get("风险维度", {})
    return_cfg = rules.get("收益维度", {})
    module_mapping = rules.get("六模块指标映射", {})

    return jsonify({
        "riskWeights": risk_cfg.get("权重", {}),
        "returnWeights": return_cfg.get("权重", {}),
        "riskMapping": module_mapping.get("风险维度", {}),
        "returnMapping": module_mapping.get("收益维度", {}),
        "gridStrategies": rules.get("九宫格处置策略", {}),
        "cutThresholds": {
            "riskLow": risk_cfg.get("分箱阈值", {}).get("低风险上限", 1.6),
            "riskHigh": risk_cfg.get("分箱阈值", {}).get("中风险上限", 2.4),
            "returnLow": return_cfg.get("分箱阈值", {}).get("低收益上限", 1.6),
            "returnHigh": return_cfg.get("分箱阈值", {}).get("中收益上限", 2.4),
        },
        "_meta": {"source": "config/discrete_rules.json",
                  "message": "前端所有What-If计算必须以此响应中的权重和阈值为唯一数据源"}
    })''')

    add_para(doc, '')
    add_para(doc, '前端消费方式（discrete.html）：', bold=True)
    add_code(doc, '''// 页面初始化时拉取规则引擎配置（唯一数据源）
let FORMULA_CONFIG = null;

async function loadFormulaConfig() {
    const resp = await fetch('/api/discrete/config');
    FORMULA_CONFIG = await resp.json();
    console.log('[What-If] 规则引擎配置已加载，来源:', FORMULA_CONFIG._meta.source);
}

// What-If计算引擎 —— 所有参数从FORMULA_CONFIG动态读取
function calcRiskSubDim(moduleIndicators, dimName) {
    const mapping = FORMULA_CONFIG.riskMapping[dimName];
    let value = 0;
    mapping.indicators.forEach(([indicatorName, weight]) => {
        value += (moduleIndicators['模块' + mapping.module][indicatorName] || 0) * weight;
    });
    if (mapping.direction === 'inverse') value = 1 - value;
    const [lo, hi] = mapping.thresholds;
    return value >= hi ? 1 : value >= lo ? 2 : 3;
}

function recalc(moduleIndicators) {
    const rw = FORMULA_CONFIG.riskWeights;
    const ew = FORMULA_CONFIG.returnWeights;
    const R_raw = calcRiskSubDim(moduleIndicators,'region') * rw.region
                + calcRiskSubDim(moduleIndicators,'contract') * rw.contract
                + calcRiskSubDim(moduleIndicators,'customer') * rw.customer
                + calcRiskSubDim(moduleIndicators,'capital') * rw.capital
                + calcRiskSubDim(moduleIndicators,'perf') * rw.perf;
    const totalR = Object.values(rw).reduce((a,b)=>a+b, 0);
    const R = R_raw / totalR;
    // E同理...
    const ct = FORMULA_CONFIG.cutThresholds;
    return { R_level: R <= ct.riskLow ? 1 : R <= ct.riskHigh ? 2 : 3, ... };
}''')

    add_para(doc, '')
    add_para(doc, '关键约束：', bold=True)
    add_bullet(doc, 'FORMULA_CONFIG 是前端唯一数据源。禁止在JS中存在任何硬编码的权重数字。')
    add_bullet(doc, '每次页面加载时重新拉取 /api/discrete/config，确保与后端discrete_rules.json即时同步。')
    add_bullet(doc, '如果API不可用，What-If面板应显示"规则配置加载失败，沙盘推演不可用"并禁用所有滑块——绝不允许使用fallback硬编码值。')

    doc.add_page_break()

    # ═══════════════ CHAPTER 5: 置信度 ═══════════════
    add_heading(doc, '第五章 数据置信度视觉反馈', 1)
    add_heading(doc, '5.1 设计理念', 2)
    add_para(doc, '模块六得分不仅作为综合评分的10%权重，更作为九宫格的"数据置信度"标识。若某项目落在"(1,3) 扩张区"但模块六得分极低（存在中标签约金额偏差、关键字段缺失），则直观提示："该项目看似高收益低风险，但底层数据严重失真，决策需谨慎"。')

    add_heading(doc, '5.2 置信度等级', 2)
    add_table(doc, ['等级','模块六得分','视觉编码','决策含义'],
        [['高','≥80','实心圆点，白色边框，opacity 0.92','数据可靠，可据此决策'],
         ['中','50-79','半透明(opacity 0.70)，橙色边框','数据存在缺失，建议复核'],
         ['低','<50','高度透明(opacity 0.45)，红色虚线3px边框，菱形符号','⚠ 数据严重失真，必须先复核底层数据']],
        [2.5,2.5,5.5,5.5])

    add_heading(doc, '5.3 ECharts编码', 2)
    add_code(doc, '''itemStyle: {
    color: confidence >= 80 ? row.color : confidence >= 50 ? row.color+'99' : row.color+'55',
    borderColor: confidence < 50 ? '#d94d63' : confidence < 80 ? '#f59e0b' : '#fff',
    borderWidth: confidence < 50 ? 3 : confidence < 80 ? 2 : 1.5,
    borderType: confidence < 50 ? 'dashed' : 'solid',
    opacity: confidence >= 80 ? 0.92 : confidence >= 50 ? 0.70 : 0.45,
},
symbol: confidence < 30 ? 'diamond' : 'circle',''')

    add_para(doc, '')
    add_para(doc, 'Tooltip增强：低置信度气泡悬停时红色警告"⚠ 数据置信度低(NN分)"，并列出该项目的具体数据质量问题（来自模型1.4输出）。散点图底部新增置信度图例条。')

    doc.add_page_break()

    # ═══════════════ CHAPTER 6: What-If ═══════════════
    add_heading(doc, '第六章 What-If沙盘推演模式', 1)
    add_heading(doc, '6.1 设计理念', 2)
    add_para(doc, 'R/S已完全公式化（六模块指标→加权→离散→九宫格），且权重由 /api/discrete/config 动态下发。前端JS根据下发的规则动态构建计算引擎，业务人员拖动滑块微调模块指标，实时观察气泡在九宫格中的位置变化——无需重新运行后端Python模型。')

    add_heading(doc, '6.2 UI布局', 2)
    add_para(doc, '在discrete.html右侧新增可折叠侧边栏"📊 沙盘推演"：')
    add_bullet(doc, '推演对象选择器（下拉框，默认当前选中项目）')
    add_bullet(doc, '六模块关键指标滑块（每模块1-2个可调节指标，滑块背景色按当前分区着色）')
    add_bullet(doc, '推演结果面板（当前定位 → 推演后定位，R/E变化量，分区跃迁方向）')
    add_bullet(doc, '[重置] [应用推演] 按钮')
    add_para(doc, '')
    add_para(doc, '推演轨迹动画：在散点图上绘制从原始位置到推演位置的绿色虚线箭头，推演气泡使用菱形+绿色边框，使用ECharts lines系列+effect动画。')

    add_heading(doc, '6.3 JS公式引擎（动态构建版）', 2)
    add_para(doc, '所有权重和阈值从 FORMULA_CONFIG（/api/discrete/config响应）动态读取，不硬编码任何数值。详见第四章代码示例。')

    doc.add_page_break()

    # ═══════════════ CHAPTER 7: 前端改造 ═══════════════
    add_heading(doc, '第七章 前端页面改造方案', 1)
    add_heading(doc, '7.1 business.html：区域-客户-履约健康监测台', 2)
    add_para(doc, '三区联动布局：顶部KPI概览行 → 中部（六轴雷达图+态势总览）→ 下部（六模块详情展开区+三级对标表）。')

    add_heading(doc, '7.1.1 异步加载与骨架屏', 3)
    add_para(doc, '后端新增3类轻量API：')
    add_bullet(doc, 'GET /api/business/modules-summary —— 6模块得分+综合评分（<1KB，秒返）')
    add_bullet(doc, 'GET /api/business/module/<id> —— 单个模块完整指标（id=1~6）')
    add_bullet(doc, 'GET /api/business/benchmark/<scope> —— 对标数据（subsidiaries/cities/projects）')
    add_para(doc, '')
    add_para(doc, '加载时序：')
    add_bullet(doc, '[0ms] 渲染6个shimmer骨架卡片 + KPI占位')
    add_bullet(doc, '[50ms] fetch modules-summary → KPI亮起 + 雷达图渲染')
    add_bullet(doc, '[50ms] Promise.allSettled并行6个module API → 哪个先返回就亮起哪个卡片')
    add_bullet(doc, '[200ms] fetch benchmark → 对标表渲染')

    add_heading(doc, '7.2 毛玻璃视觉质感升级', 2)
    add_alert(doc, '★ 全站UI升级：毛玻璃设计系统（Glassmorphism Design System）', 'info')

    add_para(doc, '为了配合整体监控台的高级感，在卡片、悬浮层及九宫格的UI上大面积应用毛玻璃效果。利用背板模糊（backdrop-filter: blur()）结合微透明的纯色背景，让复杂的数据指标"悬浮"于界面之上，极大减轻用户的视觉负担。')

    add_heading(doc, '7.2.1 设计令牌（CSS变量扩展）', 3)
    add_code(doc, ''':root {
    /* 毛玻璃层级 */
    --glass-bg: rgba(255, 255, 255, 0.72);
    --glass-bg-deep: rgba(255, 255, 255, 0.88);
    --glass-bg-light: rgba(255, 255, 255, 0.55);
    --glass-border: rgba(255, 255, 255, 0.45);
    --glass-shadow: 0 8px 32px rgba(15, 41, 76, 0.06);
    --glass-blur: blur(14px);
    --glass-blur-heavy: blur(20px);

    /* 卡片层级 */
    --card-bg: var(--glass-bg);
    --card-border: 1px solid rgba(191, 219, 254, 0.35);
    --card-radius: 16px;
    --card-shadow: 0 8px 32px rgba(15, 41, 76, 0.04);

    /* 悬浮层 */
    --overlay-bg: rgba(255, 255, 255, 0.85);
    --overlay-blur: blur(16px);

    /* 色彩语义（微透明） */
    --accent-glass: rgba(31, 137, 223, 0.08);
    --danger-glass: rgba(217, 77, 99, 0.06);
    --success-glass: rgba(35, 135, 111, 0.06);
    --warning-glass: rgba(183, 123, 18, 0.06);
}''')

    add_heading(doc, '7.2.2 核心组件CSS', 3)
    add_code(doc, '''/* 卡片基类 */
.glass-card {
    background: var(--glass-bg);
    backdrop-filter: var(--glass-blur);
    -webkit-backdrop-filter: var(--glass-blur);
    border: var(--card-border);
    border-radius: var(--card-radius);
    box-shadow: var(--card-shadow);
    transition: all 0.3s ease;
}
.glass-card:hover {
    background: var(--glass-bg-deep);
    box-shadow: 0 12px 40px rgba(15, 41, 76, 0.08);
    transform: translateY(-1px);
}

/* 模块得分卡片（六模块卡片区） */
.module-score-card {
    background: var(--glass-bg);
    backdrop-filter: var(--glass-blur);
    border: 1px solid rgba(191, 219, 254, 0.30);
    border-radius: 14px;
    padding: 18px 20px;
    position: relative;
    overflow: hidden;
}
.module-score-card::before {
    content: '';
    position: absolute;
    top: 0; left: 0; right: 0;
    height: 3px;
    border-radius: 3px 3px 0 0;
    background: linear-gradient(90deg, var(--accent-color, #1f89df), transparent);
    opacity: 0.6;
}
.module-score-card.module-strong::before { background: linear-gradient(90deg, #23876f, transparent); }
.module-score-card.module-pressure::before { background: linear-gradient(90deg, #d94d63, transparent); }

/* KPI概览卡片 */
.kpi-glass-card {
    background: rgba(255,255,255,0.65);
    backdrop-filter: blur(12px);
    border: 1px solid rgba(255,255,255,0.5);
    border-radius: 14px;
    padding: 16px 20px;
    box-shadow: 0 4px 24px rgba(15,41,76,0.03);
}

/* 九宫格气泡Tooltip */
.scatter-tooltip {
    background: var(--overlay-bg) !important;
    backdrop-filter: var(--overlay-blur) !important;
    border: 1px solid rgba(220, 232, 246, 0.6) !important;
    border-radius: 12px !important;
    box-shadow: 0 8px 32px rgba(15,41,76,0.08) !important;
}

/* What-If侧边栏 */
.whatif-panel {
    background: rgba(255,255,255,0.78);
    backdrop-filter: blur(16px);
    border-left: 1px solid rgba(220,232,246,0.5);
    box-shadow: -4px 0 24px rgba(15,41,76,0.04);
}

/* 骨架屏（毛玻璃版本） */
.skeleton-glass {
    background: rgba(232, 237, 242, 0.50);
    backdrop-filter: blur(4px);
    border-radius: 14px;
    overflow: hidden;
}
.skeleton-glass::after {
    content: '';
    position: absolute;
    top: 0; left: -100%; width: 100%; height: 100%;
    background: linear-gradient(90deg,
        transparent 0%,
        rgba(255,255,255,0.45) 50%,
        transparent 100%);
    animation: shimmer 1.8s ease-in-out infinite;
}
@keyframes shimmer {
    0% { left: -100%; }
    100% { left: 100%; }
}

/* 平滑过渡 */
.module-card-enter {
    animation: cardReveal 0.45s cubic-bezier(0.22, 0.61, 0.36, 1);
}
@keyframes cardReveal {
    from { opacity: 0; transform: translateY(12px); filter: blur(2px); }
    to   { opacity: 1; transform: translateY(0);    filter: blur(0);   }
}''')

    add_heading(doc, '7.2.3 实施范围', 3)
    add_bullet(doc, 'base.html：新增CSS变量 + .glass-card / .skeleton-glass / @keyframes cardReveal')
    add_bullet(doc, 'business.html：六模块卡片、KPI概览卡片全部采用 .module-score-card + .kpi-glass-card')
    add_bullet(doc, 'discrete.html：九宫格矩阵卡片、侧边栏、Tooltip全部采用毛玻璃样式')
    add_bullet(doc, 'dashboard.html：文件上传面板、模型卡片增加微弱毛玻璃效果（保持功能页面的克制感）')

    add_para(doc, '')
    add_para(doc, '毛玻璃设计在Chrome/Edge/Safari上完整支持。Firefox需要fallback：@supports not (backdrop-filter: blur(1px)) { .glass-card { background: rgba(255,255,255,0.94); } }', size=10)

    doc.add_page_break()

    # ═══════════════ CHAPTER 8: 内存释放 ═══════════════
    add_heading(doc, '第八章 内存释放与性能优化', 1)
    add_heading(doc, '8.1 main.py分阶段释放', 2)
    add_para(doc, 'pipeline各阶段结束后显式释放大对象：')
    add_bullet(doc, '数据加载完成 → del dmp_raw, bid_report → gc.collect()')
    add_bullet(doc, '11模型完成 → del appendices → gc.collect()')
    add_bullet(doc, '六模块完成 → del dmp, region_auth → gc.collect()（释放173字段大宽表，可能数百MB）')
    add_bullet(doc, '报告生成仅消费已缓存的pickle小对象')

    add_heading(doc, '8.2 web_app.py API级优化', 2)
    add_bullet(doc, '所有数据API响应后调用 gc.collect(0)（仅回收年轻代，开销极小）')
    add_bullet(doc, '分模块API每次只从pickle中提取单个模块数据')
    add_bullet(doc, '新增 GET /api/memory-status 运维端点（RSS/VMS/GC计数）')

    doc.add_page_break()

    # ═══════════════ CHAPTER 9: 异步并发导出对齐 ═══════════════
    add_heading(doc, '第九章 异步并发下的导出对齐防线', 1)

    add_alert(doc, '⚠ 并发陷阱：如果系统在执行异步计算（Promise.allSettled）与垃圾回收（gc.collect()）的间隙被触发底层Excel报告生成，由于内存指针与数据帧长度的瞬间波动，极易引发写入错位（窜行）。', 'danger')

    add_heading(doc, '9.1 根因分析', 2)
    add_para(doc, 'excel_exporter.py（以及相关的excel_beautifier.py）在写入Excel时依赖DataFrame的行号索引。当以下条件同时满足时，会发生窜行：')
    add_bullet(doc, '（A）前端异步请求触发了后端数据刷新（business_results被重新计算）')
    add_bullet(doc, '（B）主线程在gc.collect()中移动了内存对象')
    add_bullet(doc, '（C）Excel写入循环使用动态 len(df) 而非锁定行号')
    add_para(doc, '这三个条件在某些高并发场景下可能同时命中，导致报告的某一行数据实际属于另一个项目。')

    add_heading(doc, '9.2 补强方案：严格行号锁定', 2)
    add_para(doc, '在 excel_exporter.py（及 excel_beautifier.py）中新增以下防护：')

    add_para(doc, '')
    add_para(doc, '（1）深拷贝快照（Snapshot）', bold=True)
    add_code(doc, '''# excel_exporter.py —— 导出前创建数据快照
def export_to_excel(all_results, chain_results, output_dir, ...):
    """导出前先深拷贝全部数据，隔离后续GC影响"""
    import copy
    _all_results = copy.deepcopy(all_results)
    _chain_results = copy.deepcopy(chain_results)
    # 之后的写入操作全部使用 _all_results / _chain_results
    ...
    # 写入完成后显式释放快照
    del _all_results, _chain_results
    gc.collect(0)''')

    add_para(doc, '')
    add_para(doc, '（2）静态行号索引', bold=True)
    add_code(doc, '''# excel_beautifier.py —— 禁止动态len(df)，改为锁定行号
def write_table_with_locked_index(ws, start_row, df_snapshot, columns):
    """使用深拷贝快照+锁定行号写入，防止窜行"""
    row_count = len(df_snapshot)  # 写入前一次性确定行数
    row_indices = list(range(start_row, start_row + row_count))

    # 锁死列宽和行高
    for col_idx, col_name in enumerate(columns):
        ws.cell(row=start_row-1, column=col_idx+1, value=col_name)

    for i, (row_idx, (_, record)) in enumerate(zip(row_indices, df_snapshot.iterrows())):
        for col_idx, col_name in enumerate(columns):
            value = record.get(col_name, "")
            ws.cell(row=row_idx, column=col_idx+1, value=value)

    return start_row + row_count  # 返回下一个可用行号''')

    add_para(doc, '')
    add_para(doc, '（3）导出互斥锁', bold=True)
    add_code(doc, '''# web_app.py —— 导出API与数据刷新API互斥
export_lock = threading.Lock()

@app.route("/export/report")
def export_report():
    with export_lock:
        # 在此期间，/api/run 和 /api/run-prefilter 返回 409
        return send_file(REPORT_PATH, ...)

@app.route("/api/run", methods=["POST"])
def api_run():
    if export_lock.locked():
        return jsonify({"error": "报告导出进行中，请稍后再试"}), 409
    ...''')

    add_heading(doc, '9.3 防御层级总结', 2)
    add_table(doc, ['防御层','措施','防护目标'],
        [['L1 快照隔离','深拷贝所有数据后再写入Excel','防止GC移动内存导致指针错位'],
         ['L2 静态索引','锁定row_indices=list(range(...))，禁止动态len(df)','防止DataFrame长度变化导致窜行'],
         ['L3 互斥锁','export_lock确保导出与数据刷新互斥','防止并发写入导致文件损坏'],
         ['L4 写入验证','导出后立即重新读取Excel，校验行数=快照行数','事后检测，发现窜行则自动重试一次']],
        [2.5,8.0,5.5])

    doc.add_page_break()

    # ═══════════════ CHAPTER 10: 数据容错防线 ═══════════════
    add_heading(doc, '第十章 数据容错边界与抵消逻辑防线', 1)

    add_alert(doc, '⚠ 关键数据质量风险：DMP系统深度融合了全量物资数据。在Module 4（履约盈利健康度）进行上卷聚合时，材料抵消与退场核销逻辑会产生负向产值数据。若聚合层不当使用abs()函数，将导致成本双倍扣减、模块得分断崖式失真。', 'danger')

    add_heading(doc, '10.1 问题场景', 2)
    add_para(doc, 'DMP系统的"实际完成产值"字段中包含以下复杂业务场景的数据：')
    add_bullet(doc, '场景A：材料调拨抵消。项目A向项目B调拨钢材，项目A产值减去调出金额（负值），项目B产值加上调入金额（正值）。全量聚合时应正负相抵=0。')
    add_bullet(doc, '场景B：退场核销。分包商退场，已报未完工的产值需核销为负数。这是真实的业务损失，必须保留负向差异。')
    add_bullet(doc, '场景C：甲方扣款。质量问题导致的甲方扣款，产值调减。同样必须保留负向差异。')
    add_para(doc, '如果在Module 4聚合时错误地对产值字段取绝对值（abs()），场景B和C的损失将被抹去，导致盈利健康度假性偏高。')

    add_heading(doc, '10.2 补强方案：分层abs()策略', 2)
    add_para(doc, '将abs()防护限定在数据清洗层（model输入层），禁止在聚合层（module计算层）再次调用abs()：')

    add_table(doc, ['层级','abs()策略','说明'],
        [['L0 原始数据层','不使用abs()','读取DMP原始值，保留正负号'],
         ['L1 清洗层（data_adapter.py）','仅对明确的录入错误取abs()\n（如签约额=-100的情况）','窄范围、白名单式清洗，\n不触碰产值/收款字段'],
         ['L2 模型层（11个model）','不使用abs()','模型直接使用清洗后的含正负号数据'],
         ['L3 六模块聚合层','★ 严禁使用abs() ★','直接Σ求和，正负自然相抵。\n场景A自动抵消，场景B/C保留真实损失'],
         ['L4 前端展示层','取abs()仅用于百分比计算的分母\n（如产值转化率的分母取abs(签约额)）','分子（产值）保留符号，\n分母（签约额）取绝对值防负号干扰']],
        [2.5,5.5,8.0])

    add_heading(doc, '10.3 代码级防护', 2)
    add_code(doc, '''# business_analysis.py —— _performance_module() 严禁abs
def _performance_module(self, group, ...):
    """模块四：履约盈利健康度 —— 产值聚合严禁abs()"""
    # ✅ 正确：直接求和，正负自然相抵
    total_output = group["_actual_output"].sum()

    # ❌ 错误（已封禁）：total_output = group["_actual_output"].abs().sum()
    # 如果检测到abs()调用，在CI/CR中直接拦截

    # 产值转化率：分母取abs防负号，分子保留符号
    total_contract = group["_contract_amt"].abs().sum()  # 分母abs安全
    conversion_rate = total_output / total_contract if total_contract > 0 else 0.0
    # 注意：如果total_output为负（材料调出>调入），conversion_rate可以为负，
    # 此时在前端展示为"0%"并附加提示"该单位/城市本期材料净调出"
    ...

# 增加运行时断言
assert not any(call.get('abs') for call in traceback.format_stack()
               if '_actual_output' in str(call)), \
    "CRITICAL: abs() detected on _actual_output in aggregation layer!"''')

    add_heading(doc, '10.4 验证机制', 2)
    add_bullet(doc, '单元测试：构造场景A数据（正负相抵），验证Module 4聚合结果=0而非>0')
    add_bullet(doc, '单元测试：构造场景B数据（退场核销负值），验证负值被保留而非被abs()抹去')
    add_bullet(doc, 'CI Lint规则：grep检测business_analysis.py中是否存在 _actual_output.*\\.abs\\(\\)，命中则构建失败')

    doc.add_page_break()

    # ═══════════════ CHAPTER 11: 实施步骤 ═══════════════
    add_heading(doc, '第十一章 实施步骤与工时估算', 1)
    add_table(doc, ['阶段','核心工作','工时'],
        [['P1 六模块补全','business_analysis.py新增模块六+扩展全指标+web_app.py新增分模块API', '4.5h'],
         ['P2 规则引擎双源','web_app.py新增/api/discrete/config + discrete.html重构JS为动态公式引擎', '2h'],
         ['P3 监测台重构','business.html骨架屏+异步加载+六模块卡片+毛玻璃样式+雷达图+对标表', '6h'],
         ['P4 R/S映射','discrete_rules.json新增映射段+discrete_analysis.py新增run_with_module_scores+main.py调整', '3.5h'],
         ['P5 置信度','置信度计算+散点图视觉编码', '2h'],
         ['P6 What-If','基于/api/discrete/config动态构建的沙盘推演+滑块UI+推演动画', '3h'],
         ['P7 数据容错','Module 4 abs()防线+单元测试+Lint规则', '1.5h'],
         ['P8 导出对齐','excel_exporter深拷贝快照+静态行号索引+导出互斥锁', '1.5h'],
         ['P9 联调','端到端测试+边界处理+Firefox毛玻璃fallback验证', '2h']],
        [3.0,11.0,2.0])
    add_para(doc, '')
    add_para(doc, '总计工时：约26小时', bold=True)

    add_heading(doc, '11.1 数据流顺序（main.py调整后）', 2)
    add_para(doc, '[1]数据加载 → [2]前置过滤 → [3]11模型运行 → [4]三条链 → [5]六模块分析(新增) → [6]九宫格离散分析(消费六模块结果) → [7]报告生成')

    add_heading(doc, '11.2 向后兼容', 2)
    add_para(doc, '如果六模块分析尚未执行（无business_results缓存），discrete_analysis.py自动回退到v2.9的run()方法。如果/api/discrete/config不可用，What-If面板禁用并提示——绝不允许fallback硬编码值。')

    doc.add_page_break()

    # ═══════════════ CHAPTER 12: 文件清单 ═══════════════
    add_heading(doc, '第十二章 涉及文件清单', 1)
    add_table(doc, ['文件','改动类型','改动说明'],
        [['models/business_analysis.py','★重写','新增模块六；扩展全指标；abs()防线；聚合层签名验证'],
         ['templates/business.html','★重写','三区布局；骨架屏；六轴雷达；六模块卡片；毛玻璃样式'],
         ['models/discrete_analysis.py','★重写','run_with_module_scores()；置信度计算；回退兼容'],
         ['config/discrete_rules.json','●扩展','新增"六模块指标映射"配置段'],
         ['templates/discrete.html','●扩展','动态公式引擎(/api/discrete/config)；置信度编码；What-If面板；毛玻璃样式'],
         ['web_app.py','●扩展','/api/discrete/config；3个分模块API；/api/memory-status；export_lock；cleanup装饰器'],
         ['main.py','●调整','步骤顺序；分阶段gc；传入business_results'],
         ['report/excel_exporter.py','●加固','深拷贝快照+静态行号索引'],
         ['report/excel_beautifier.py','●加固','锁定行号写入+禁止动态len(df)'],
         ['templates/base.html','○微调','毛玻璃CSS变量+shimmer骨架屏+cardReveal动画']],
        [4.5,2.0,9.5])
    add_para(doc, '')
    add_para(doc, '图例：★重写(>200行)  ●扩展(50-200行)/加固  ○微调(<50行)', bold=True)

    doc.add_page_break()

    # ═══════════════ APPENDIX ═══════════════
    add_heading(doc, '附录A：核心代码框架', 1)
    add_heading(doc, 'A.1 business_analysis.py —— 模块六新增', 2)
    add_code(doc, '''def _data_quality_module(self, group, project_codes, issue_index, total_projects):
    """模块六：数据质量与流程效率分析（基于模型1.4输出）"""
    r14_df = self._get_model_output("1.4")
    key_fields = ["项目编码","项目名称","申报单位","项目地址",
                  "签约额（元）","客户名称","工程类别","签约时间"]
    completeness = sum(group[col].notna().mean() for col in key_fields
                       if col in group.columns) / len(key_fields)
    process_ok = 1.0 - self._ratio_in_issues(r14_df, project_codes, "流程", "时序")
    bid_dev_ratio = self._calc_bid_deviation_ratio(group)
    a_std = group["_a_value"].std()
    estimation_regularity = max(0, 1 - a_std / 0.05)
    sign_delay_ratio = self._calc_sign_delay_ratio(group)
    score = (completeness*0.20 + process_ok*0.25 + (1-bid_dev_ratio)*0.20
             + estimation_regularity*0.15 + (1-sign_delay_ratio)*0.20) * 100
    return ModuleScore(score=max(0,min(100,score)),
        metrics={"数据完整率":completeness,"流程合规率":process_ok,
                 "中标签约偏差率":bid_dev_ratio,
                 "测算规律性指数":estimation_regularity,
                 "签约延迟率":sign_delay_ratio})''')

    add_heading(doc, 'A.2 discrete_analysis.py —— 基于六模块的R/E计算', 2)
    add_code(doc, '''def run_with_module_scores(self, all_results, dmp_df, business_results, ...):
    module_index = self._build_module_index(business_results)
    for idx, row in dmp_df.iterrows():
        unit = str(row.get("申报单位", ""))
        unit_mod = module_index.get(unit, self._default_modules())
        # R从六模块指标计算
        r_region   = self._module_indicator_to_level(unit_mod, "跨区域经营指数", "inverse")
        r_contract = self._module_indicator_to_level(unit_mod,
                        [("风险项目占比",0.6),("合同条款不利度",0.4)], "inverse")
        r_customer = self._module_indicator_to_level(unit_mod,
                        [("客户稳定性指数",0.5),("客户集中度风险",0.5)], "inverse")
        r_capital  = self._module_indicator_to_level(unit_mod,
                        [("逾期回收率",0.6),("负流项目占比",0.4)], "inverse")
        r_perf     = self._module_indicator_to_level(unit_mod,
                        [("停工退场率",0.6),("签约履约偏差率",0.4)], "inverse")
        # E从六模块指标计算
        e_profit     = self._module_indicator_to_level(unit_mod, "盈利健康度", "direct")
        e_conversion = self._module_indicator_to_level(unit_mod, "产值转化率", "direct")
        e_collection = self._module_indicator_to_level(unit_mod, "资金回收率", "direct")
        e_scale      = self._module_indicator_to_level(unit_mod, "区域合同额强度", "direct")
        # 加权+归一化+分箱（公式不变）
        ...
        confidence = unit_mod.get("模块六",{}).get("score",50)''')

    add_heading(doc, 'A.3 web_app.py —— /api/discrete/config 端点', 2)
    add_code(doc, '''@app.route("/api/discrete/config")
def api_discrete_config():
    """返回完整的离散化规则配置，前端What-If唯一数据源"""
    disc_path = PROJECT_ROOT / "config" / "discrete_rules.json"
    with open(disc_path, "r", encoding="utf-8") as f:
        raw = json.load(f)
    rules = raw.get("离散化分析", {})
    risk_cfg = rules.get("风险维度", {})
    return_cfg = rules.get("收益维度", {})
    module_mapping = rules.get("六模块指标映射", {})
    return jsonify({
        "riskWeights": risk_cfg.get("权重", {}),
        "returnWeights": return_cfg.get("权重", {}),
        "riskMapping": module_mapping.get("风险维度", {}),
        "returnMapping": module_mapping.get("收益维度", {}),
        "gridStrategies": rules.get("九宫格处置策略", {}),
        "cutThresholds": {
            "riskLow": risk_cfg.get("分箱阈值",{}).get("低风险上限",1.6),
            "riskHigh": risk_cfg.get("分箱阈值",{}).get("中风险上限",2.4),
            "returnLow": return_cfg.get("分箱阈值",{}).get("低收益上限",1.6),
            "returnHigh": return_cfg.get("分箱阈值",{}).get("中收益上限",2.4),
        },
        "_meta": {"source":"config/discrete_rules.json",
                  "message":"前端所有What-If计算必须以此响应为唯一数据源"}
    })''')

    add_heading(doc, 'A.4 discrete.html —— 动态公式引擎', 2)
    add_code(doc, '''let FORMULA_CONFIG = null;  // ← 唯一数据源，禁止硬编码

async function loadFormulaConfig() {
    FORMULA_CONFIG = await (await fetch('/api/discrete/config')).json();
    if (!FORMULA_CONFIG || !FORMULA_CONFIG.riskWeights) {
        document.getElementById('whatifPanel').innerHTML =
            '<div class="error">规则配置加载失败，沙盘推演不可用</div>';
        return false;
    }
    return true;
}

function calcRiskSubDim(moduleIndicators, dimName) {
    const mapping = FORMULA_CONFIG.riskMapping[dimName];
    if (!mapping) return 2;
    let value = 0;
    mapping.indicators.forEach(([indicatorName, weight]) => {
        value += (moduleIndicators['模块'+mapping.module][indicatorName] || 0) * weight;
    });
    if (mapping.direction === 'inverse') value = 1 - value;
    const [lo, hi] = mapping.thresholds;
    return value >= hi ? 1 : value >= lo ? 2 : 3;
}

function recalc(moduleIndicators) {
    if (!FORMULA_CONFIG) return null;
    const rw = FORMULA_CONFIG.riskWeights;
    const ew = FORMULA_CONFIG.returnWeights;
    const dims = ['region','contract','customer','capital','perf'];
    let R_raw = 0;
    dims.forEach(d => R_raw += calcRiskSubDim(moduleIndicators, d) * (rw[d]||0));
    const totalR = Object.values(rw).reduce((a,b)=>a+b, 0);
    const R = R_raw / totalR;
    // E同理...
    const ct = FORMULA_CONFIG.cutThresholds;
    return {
        R: Math.max(1, Math.min(3, R)),
        R_level: R <= ct.riskLow ? 1 : R <= ct.riskHigh ? 2 : 3,
        E: /* ... */, E_level: /* ... */
    };
}''')

    add_heading(doc, 'A.5 excel_exporter.py —— 导出对齐防线', 2)
    add_code(doc, '''import copy, threading

export_lock = threading.Lock()

def export_to_excel(all_results, chain_results, output_dir, ...):
    """导出前深拷贝快照 + 锁定行号"""
    with export_lock:
        _all = copy.deepcopy(all_results)
        _chain = copy.deepcopy(chain_results)

        # 锁定行号索引
        row_index = 1
        for model_id, (df, summary) in _all.items():
            if df is None or len(df) == 0: continue
            row_count = len(df)  # 一次性确定
            row_indices = range(row_index, row_index + row_count)
            for i, (_, record) in enumerate(df.iterrows()):
                ws.cell(row=row_indices[i], column=1, value=record.get("项目编码",""))
            row_index += row_count

        del _all, _chain
        gc.collect(0)
    return output_path''')

    # ── Save ──
    path = os.path.join(DOCS_DIR, '实施方案_六模块九宫格全链路贯通_v2.10.docx')
    doc.save(path)
    print(f'[OK] Standalone: {path}')
    return path


# ═══════════════════════════════════════════════════════════
# DOCUMENT 2: v2.10 UPGRADE
# ═══════════════════════════════════════════════════════════
def update_v29_document():
    src = os.path.join(DOCS_DIR, '方案_v2.9_模型2.5签约履约偏差_完整版修订.docx')
    import zipfile
    z = zipfile.ZipFile(src)
    import xml.etree.ElementTree as ET
    xml_content = z.read('word/document.xml')
    tree = ET.fromstring(xml_content)
    paras = []
    for p in tree.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
        texts = []
        for t in p.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
            if t.text: texts.append(t.text)
        if texts: paras.append(''.join(texts))
    existing = '\n'.join(paras)

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)

    note = doc.add_paragraph()
    run = note.add_run('【v2.10增量更新】本文档为v2.9方案基础上新增"第八章：六模块·九宫格·全链路贯通实施方案（含双源一致性、数据容错、导出对齐、毛玻璃质感升级）"。')
    run.font.name = '微软雅黑'; run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11); run.bold = True; run.font.color.rgb = RGBColor(0x1f,0x89,0xdf)
    note2 = doc.add_paragraph()
    run2 = note2.add_run(f'更新日期：{datetime.date.today().isoformat()}。原有第一章~第七章及附录A~F完整保留于原始docx文件。新增第八章为本次增量。')
    run2.font.name = '微软雅黑'; run2._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run2.font.size = Pt(10)

    doc.add_paragraph()
    add_heading(doc, '（原有内容：第一章~第七章、附录A~F，完整保留于原始docx文件中。共'+f'{len(existing):,}字符、{existing.count(chr(10)):,}行。）', 2)

    doc.add_page_break()

    # ════════ NEW CHAPTER 8 ════════
    add_heading(doc, '第八章（v2.10新增）：六模块·九宫格·全链路贯通实施方案', 1)
    add_para(doc, '本章为v2.10核心增量更新，目标："11模型→三维度→三条链→六模块→九宫格"全链路贯通。', bold=True)

    add_heading(doc, '8.1 背景与问题诊断', 2)
    for i, p in enumerate([
        '断层一：六模块不完整。仅5模块，缺"模块六：数据质量与流程效率分析"。',
        '断层二：指标展示不全。business.html仅展示6个零散KPI。',
        '断层三：九宫格R/S与六模块脱节。DiscreteAnalyzer绕过六模块直接读模型输出。',
        '断层四：缺少置信度反馈。数据质量差的扩张区项目无法视觉区分。',
        '断层五：缺少推演能力。无法模拟指标改善后的九宫格移动。',
        '断层六（关键）：前端规则引擎若硬编码JS权重，与后端discrete_rules.json不同步时将撕裂。',
    ], 1):
        add_para(doc, f'({i}) {p}')

    add_heading(doc, '8.2 六模块完整指标体系', 2)
    add_table(doc, ['模块','对应模型','核心指标'],
        [['模块一：区域布局健康度','1.1+1.2','区域渗透率/跨区域经营指数/深耕区域集中度/区域合同额强度/业务结构偏离度/EPC转型进度'],
         ['模块二：客户资源稳定性','1.3+3.1+3.2','客户稳定性指数/客户产出波动率/客户集中度风险/中标转化率/新客户质量指数/战略客户产出比'],
         ['模块三：合同质量与风险集中度','2.1+2.4','风险项目占比/风险合同额集中度/付款条件优良率/合同条款不利度/三证合规率'],
         ['模块四：履约盈利健康度','2.2+2.5+前置过滤','产值转化率/签约履约偏差率/盈利健康度/停工退场率/效益偏差率/在施项目活跃度'],
         ['模块五：资金效率与安全性','2.3','资金占用率/保证金周转天数/逾期回收率/预收款缺口率/负流项目占比'],
         ['模块六：数据质量与流程效率','1.4','数据完整率/流程合规率/中标签约偏差率/测算规律性指数/签约延迟率']],
        [3.5,3.0,9.5])
    add_para(doc, '综合得分 = 模块一(25%)+模块二(20%)+模块三(18%)+模块四(15%)+模块五(12%)+模块六(10%)', bold=True)

    add_heading(doc, '8.3 九宫格R/S与六模块指标映射', 2)
    add_para(doc, '核心变更：九宫格R/S子维度不再直接读取11模型原始输出，改为直接采用六模块标准化指标。', bold=True)
    add_table(doc, ['R子维度','权重','六模块指标','模块'],
        [['区域合规','1.0','跨区域经营指数','模块一'],
         ['合同底线','1.0','风险项目占比(0.6)+合同条款不利度(0.4)','模块三'],
         ['客户健康','0.8','客户稳定性指数(0.5)+客户集中度风险(0.5)','模块二'],
         ['资金安全','0.8','逾期回收率(0.6)+负流项目占比(0.4)','模块五'],
         ['履约真实','0.4','停工退场率(0.6)+签约履约偏差率(0.4)','模块四']], [2.0,1.0,9.0,4.0])
    add_para(doc, '')
    add_table(doc, ['E子维度','权重','六模块指标','模块'],
        [['盈利水平','1.05','盈利健康度','模块四'],
         ['产值转化','0.75','产值转化率','模块四'],
         ['资金回收','0.75','资金回收率(收款/签约额)','模块五'],
         ['合同规模','0.45','区域合同额强度(归一化)','模块一']], [2.0,1.0,9.0,4.0])

    add_heading(doc, '8.4 规则引擎双源一致性（防What-If致命伤）', 2)
    add_alert(doc, '⚠ 架构约束：绝对禁止在JS中硬编码任何权重、阈值或映射关系。新增GET /api/discrete/config端点，前端所有What-If计算必须以此响应为唯一数据源。若API不可用，What-If面板禁用并提示——不允许fallback硬编码值。', 'danger')

    add_heading(doc, '8.5 数据置信度视觉反馈', 2)
    add_para(doc, '模块六得分作为"数据置信度"标识。低置信度(<50分)气泡：虚线红色边框、高度透明(opacity 0.45)、菱形符号。Tooltip红色警告"⚠ 数据置信度低，决策需谨慎"。')

    add_heading(doc, '8.6 What-If沙盘推演', 2)
    add_para(doc, '右侧可折叠侧边栏+滑块调节六模块指标+JS根据/api/discrete/config动态构建公式引擎+气泡实时跃迁+绿色虚线箭头轨迹动画。')

    add_heading(doc, '8.7 数据容错与抵消逻辑防线', 2)
    add_alert(doc, '⚠ Module 4聚合层严禁abs()：材料抵消(A)和退场核销(B/C)的负向差异必须保留。abs()仅限L1清洗层（窄范围白名单），L3聚合层直接Σ求和，正负自然相抵。CI Lint规则：grep检测business_analysis.py中_actual_output.*\.abs\(\) → 构建失败。', 'danger')

    add_heading(doc, '8.8 异步并发导出对齐', 2)
    add_alert(doc, '⚠ excel_exporter.py + excel_beautifier.py必须：①深拷贝快照后再写入 ②锁定行号索引(list(range))禁止动态len(df) ③export_lock互斥锁。防止GC并发下窜行。', 'warning')

    add_heading(doc, '8.9 前端异步加载与骨架屏 + 毛玻璃质感', 2)
    add_para(doc, 'business.html采用Promise.allSettled分块异步加载，骨架屏shimmer动画。全站应用毛玻璃设计系统（backdrop-filter: blur + 微透明背景），卡片悬浮感+cardReveal入场动画。')

    add_heading(doc, '8.10 内存释放', 2)
    add_para(doc, 'main.py分阶段del+gc.collect()。web_app.py API响应后gc.collect(0)。新增/api/memory-status运维端点。')

    add_heading(doc, '8.11 实施步骤', 2)
    add_table(doc, ['阶段','核心工作','工时'],
        [['P1 六模块补全','模块六+全指标+分模块API','4.5h'],
         ['P2 规则引擎双源','/api/discrete/config + 动态公式引擎','2h'],
         ['P3 监测台重构','骨架屏+异步+六模块卡片+毛玻璃','6h'],
         ['P4 R/S映射','配置+run_with_module_scores+main.py调整','3.5h'],
         ['P5 置信度','置信度计算+散点图编码','2h'],
         ['P6 What-If','动态公式引擎+滑块UI+动画','3h'],
         ['P7 数据容错','abs()防线+单元测试+Lint','1.5h'],
         ['P8 导出对齐','快照+静态索引+互斥锁','1.5h'],
         ['P9 联调','端到端+Firefox fallback','2h']],
        [3.0,11.0,2.0])
    add_para(doc, '')
    add_para(doc, '总计工时：约26小时', bold=True)

    add_heading(doc, '8.12 涉及文件清单', 2)
    add_table(doc, ['文件','改动','说明'],
        [['models/business_analysis.py','★重写','模块六+全指标+abs()防线'],
         ['templates/business.html','★重写','三区布局+骨架屏+毛玻璃'],
         ['models/discrete_analysis.py','★重写','run_with_module_scores+置信度'],
         ['config/discrete_rules.json','●扩展','六模块指标映射段'],
         ['templates/discrete.html','●扩展','动态公式引擎+置信度编码+What-If+毛玻璃'],
         ['web_app.py','●扩展','/api/discrete/config+3个分模块API+memory-status+export_lock'],
         ['main.py','●调整','步骤顺序+分阶段gc'],
         ['report/excel_exporter.py','●加固','深拷贝快照+静态行号索引'],
         ['report/excel_beautifier.py','●加固','锁定行号+禁止动态len(df)'],
         ['templates/base.html','○微调','毛玻璃CSS变量+shimmer+cardReveal']],
        [4.5,2.0,9.5])

    add_heading(doc, '8.13 向后兼容', 2)
    add_para(doc, '无business_results→回退v2.9 run()。无/api/discrete/config→What-If禁用（不允许fallback硬编码）。Module 4 CI Lint保abs()防线不被后续提交破坏。')

    from datetime import datetime as dt
    ts = dt.now().strftime('%Y%m%d_%H%M%S')
    path = os.path.join(DOCS_DIR, f'方案_v2.10_六模块九宫格全链路贯通_完整版_{ts}.docx')
    doc.save(path)
    print(f'[OK] v2.10 upgrade: {path}')
    return path


# ── Run ──
if __name__ == '__main__':
    print('='*60)
    print('  生成Word文档 (v2 + 4 supplementary sections)')
    print('='*60)
    p1 = create_standalone_plan()
    p2 = update_v29_document()
    print(f'\n生成文件：\n  1. {p1}\n  2. {p2}\nDone!')
