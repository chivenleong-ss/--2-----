"""
Update model 1.4 section in the v2.9 Word document:
- Replace outdated OA-dependent description with DMP-first/OA-second strategy
- Update field gap description
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r'C:\Users\sasa\Desktop\模型建设\模型2：市场营销\【完善版修订】全面数字化营销审计系统实施方案_v2.9.docx'

doc = Document(SRC)

# Find model 1.4 section paragraphs
m14_start = None
for i, p in enumerate(doc.paragraphs):
    if '模型1.4：营销统计数据多维度交叉验真模型' in p.text:
        m14_start = i
        break

if m14_start is None:
    print('ERROR: Model 1.4 section not found')
    sys.exit(1)

print(f'Model 1.4 section starts at P{m14_start}')

# Find key paragraphs to update
for i in range(m14_start, m14_start + 35):
    if i >= len(doc.paragraphs):
        break
    txt = doc.paragraphs[i].text

    # P368: Update DMP timeline check section
    if '【DMP时间线校验】' in txt or ('DMP时间线校验' in txt):
        doc.paragraphs[i].text = (
            '  【DMP四节点时间线校验（v2.9升级，纯DMP字段，无需OA）】'
        )
        # Next paragraph (P369) should be updated
        if i + 1 < len(doc.paragraphs):
            doc.paragraphs[i + 1].text = (
                '  1a. 领取招标文件时间 vs 交标时间 → 领取晚于交标 = 招文领取与交标时间倒置（red，疑似事后补录）\n'
                '  1b. 交标时间 vs 中标时间 → 交标晚于中标 = 交标与中标时间倒置（red，疑似事后补录）\n'
                '  1c. 中标时间 vs 签约报量时间 → 中标晚于报量 = 中标与签约报量时间倒置（yellow，时序异常）\n'
                '  1d. 签约报量时间 vs 签约时间 → 报量晚于签约 = 签约报量与签约时间倒置（yellow，报量不应晚于签约）'
            )

    # P374: Update to remove redundancy
    if '2c. 中标后>180天未签约 → 中标即无后续签约' in txt:
        doc.paragraphs[i].text = (
            '  2c. 中标后>180天未签约 → 中标即无后续签约（僵尸投标）'
        )

    # P377: Replace the old OA-dependent check description
    if '招文评审通过时间 vs 交标时间 → 未批先投违规' in txt:
        doc.paragraphs[i].text = (
            '  【DMP已覆盖，OA仅做二级验证】\n'
            '  3. 招文评审通过时间 vs 交标时间 → 精确的未批先投检测。\n'
            '     v2.9策略：DMP四节点时间线先完成全量初筛（1a-1d），\n'
            '     仅当DMP检测出时间倒置疑点时，再调取OA审批流精确时间做二级验证。\n'
            '     OA不再作为一级检测依赖，改为"疑点定向核查"模式。'
        )

    # P384-386: Update field gap description
    if '字段缺口说明' in txt:
        doc.paragraphs[i].text = '**字段缺口说明（v2.9更新）**：'

    if '中标报量接入后，全部6项校验均已完成' in txt:
        doc.paragraphs[i].text = (
            'v2.9版本，DMP四节点时间线已实现4项时序倒置检测（1a-1d），'
            '中标报量补充中标额/预计签约时间后完成中标签约交叉验证（2a-2c），'
            '加上直接发包/邀请招标比例检测（#4）、利润率规律异常（#5）、'
            '集中签约异常（#6），全部检测均可在无OA数据情况下独立运行。\n\n'
            'OA审批流时间戳的新定位：\n'
            '• 不再作为一级检测依赖 —— DMP时间线已能发现时间倒置疑点\n'
            '• 改为二级验证工具 —— 当DMP检测出时间倒置时，调取OA精确审批时间做定向核查\n'
            '• 这样做的好处：(1) 不依赖OA即可全量扫描；(2) 减少OA数据采集范围（仅需调取疑点项目）；'
            '(3) OA数据获取有明确触发条件，避免盲目全量导出'
        )

print('Model 1.4 section updated.')

# Also update P996 (Appendix C: OA审批流时间戳)
for i, p in enumerate(doc.paragraphs):
    if 'C. OA审批流时间戳' in p.text and p.style.name.startswith('Heading'):
        print(f'Found Appendix C at P{i}')
        # Read next few paragraphs
        for j in range(i+1, min(i+10, len(doc.paragraphs))):
            txt = doc.paragraphs[j].text
            if 'OA审批流' in txt or '招标文件评审' in txt or '立项审批' in txt or '合同评审' in txt:
                doc.paragraphs[j].text = (
                    '**v2.9策略调整**：OA审批流时间戳不再作为模型1.4的一级检测依赖。'
                    'DMP签约报量（四局）提供的5个时间节点（领取招标文件时间、交标时间、'
                    '中标时间、签约报量时间、签约时间）已实现4项全量时序倒置初筛。'
                    'OA审批流精确时间改为"疑点定向核查"模式：仅当DMP检测出时间倒置疑点时，'
                    '才需要调取对应项目的OA审批时间做二级验证。'
                    '此策略可大幅减少OA数据采集范围，同时不降低检测覆盖率。'
                )
                break
        break

# Update revision date
doc.paragraphs[1].text = '**修订日期**：2026-06-02'

# Update P6 to add this change
old_p6 = doc.paragraphs[6].text
if '停缓建' not in old_p6:
    doc.paragraphs[6].text = old_p6.rstrip('。）') + '；（7）模型1.4 OA策略调整为"DMP初筛→疑点定向核查OA"二级模式，OA不再作为一级检测依赖。）'

# Save to new file (original may be locked by Word)
import time
ts = time.strftime('%H%M%S')
DST = SRC.replace('_v2.9.docx', f'_v2.9_{ts}.docx')
doc.save(DST)
print(f'Saved: {DST}')
print('Done!')
