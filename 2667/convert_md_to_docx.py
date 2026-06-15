import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_shading(cell, color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def parse_table(doc, lines, i):
    rows_data = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [c.strip() for c in lines[i].strip().split("|")[1:-1]]
        rows_data.append(cells)
        i += 1

    if len(rows_data) < 2:
        return i

    header = rows_data[0]
    data = rows_data[2:]
    ncols = len(header)
    table = doc.add_table(rows=1 + len(data), cols=ncols, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, txt in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(txt)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        set_shading(cell, "D9E2F3")

    for ri, row in enumerate(data):
        for cj, txt in enumerate(row):
            if cj < ncols:
                cell = table.rows[ri + 1].cells[cj]
                cell.text = ""
                run = cell.paragraphs[0].add_run(txt)
                run.font.size = Pt(9)
                run.font.name = "Microsoft YaHei"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    doc.add_paragraph()
    return i


def run():
    md_path = r"c:\Users\sasa\Desktop\模型建设\模型2：市场营销\方案_v2.9_模型2.5签约履约偏差_代码对齐修订稿.md"
    docx_path = r"c:\Users\sasa\Desktop\模型建设\模型2：市场营销\方案_v2.9_模型2.5签约履约偏差_代码对齐修订稿.docx"

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    i = 0
    in_code = False
    code_buf = []

    while i < len(lines):
        line = lines[i].rstrip("\n\r")
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                for cl in code_buf:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.left_indent = Cm(1)
                    run = p.add_run(cl)
                    run.font.name = "Consolas"
                    run.font.size = Pt(8)
                doc.add_paragraph()
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("|"):
            i = parse_table(doc, lines, i)
            continue

        h1 = re.match(r"^# (.+)", stripped)
        h2 = re.match(r"^## (.+)", stripped)
        h3 = re.match(r"^### (.+)", stripped)

        if h1:
            heading = doc.add_heading(h1.group(1), level=1)
            for run in heading.runs:
                run.font.name = "Microsoft YaHei"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            i += 1
            continue

        if h2:
            heading = doc.add_heading(h2.group(1), level=2)
            for run in heading.runs:
                run.font.name = "Microsoft YaHei"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            i += 1
            continue

        if h3:
            heading = doc.add_heading(h3.group(1), level=3)
            for run in heading.runs:
                run.font.name = "Microsoft YaHei"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            i += 1
            continue

        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(stripped[2:])
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph("─" * 60)
            i += 1
            continue

        p = doc.add_paragraph(stripped)
        for run in p.runs:
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        i += 1

    doc.save(docx_path)
    print("OK")


if __name__ == "__main__":
    run()
