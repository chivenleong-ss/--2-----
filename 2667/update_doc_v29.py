"""
Update the existing implementation plan document to v2.9.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy
import datetime

SRC = r'C:\Users\sasa\Desktop\模型建设\模型2：市场营销\【完善版修订】全面数字化营销审计系统实施方案.docx'
DST = r'C:\Users\sasa\Desktop\模型建设\模型2：市场营销\【完善版修订】全面数字化营销审计系统实施方案_v2.9.docx'

doc = Document(SRC)

# ====================================================================
# Helper: find paragraph index by text match
# ====================================================================
def find_para(text_fragment, start=0):
    for i, p in enumerate(doc.paragraphs):
        if i >= start and text_fragment in p.text:
            return i
    return None

def find_heading(text_fragment):
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith('Heading') and text_fragment in p.text:
            return i
    return None

# ====================================================================
# 1. UPDATE TITLE & REVISION INFO
# ====================================================================
print('1. Updating title and revision info...')

# Title
doc.paragraphs[0].text = '全面数字化营销审计系统端到端落地实施方案（完善版修订稿 v2.9）'

# Revision date
doc.paragraphs[1].text = '**修订日期**：2026-06-01'

# Revision notes - update paragraph 2
doc.paragraphs[2].text = (
    '**修订说明**：本稿在完善版方案基础上，依据**中国建筑第四工程局有限公司市场营销管理手册0520'
    '（2026年5月修订版）**全面修订，并在v2.9版本中完成以下重大更新。'
)

# Add v2.9 core updates to paragraph 6
doc.paragraphs[6].text = (
    'v2.9核心更新：（1）DMP四节点时间线校验替代OA审批流（交标→中标→签约报量→签约），'
    '模型1.4新增4项时序倒置检测；（2）模型2.1新增「是否付款条件不达标」制度交叉校验，'
    '独立验证DMP标记填列正确性；（3）签约报量（四局）170字段导出适配，新增累计垫资周期、'
    '付款周期等数值转换；（4）三证字段确认DMP已含，移除App1依赖；'
    '（5）offline_fields从40字段扩展至63字段（+23个DMP字段）；'
    '（6）代码审查修复：model_2_2僵尸合同筛查变量未定义、report版本号及摘要key更新。'
)

# ====================================================================
# 2. UPDATE CHAPTER 1 - DATA SOURCE MATRIX
# ====================================================================
print('2. Updating data source description...')

# P11: Update the data source description
p11_idx = find_para('DMP线上导出数据为主表')
if p11_idx:
    doc.paragraphs[p11_idx].text = (
        '说明：DMP线上导出数据（签约报量四局170字段，v2.9确认）为主表全覆盖，'
        '审计附表（手工填列）仅补充DMP中没有的13个财务实际数据字段。'
        '中标报量导出补充投标阶段时间线及金额字段（中标额、预计签约时间等）。'
        '区域认定按年度分表——系统根据项目签约年份自动匹配对应年度的认定表'
        '（≤2023→2023，2024→2024，2025→2025，2026→2026）。'
    )

# P13: "173个字段" → "170个字段"
p13_idx = find_para('从DMP系统173个字段中')
if p13_idx:
    doc.paragraphs[p13_idx].text = (
        '从DMP系统170个字段（签约报量四局v2.9确认）中，按审计模型实际需要用到的字段梳理如下：'
    )

# P18: Update section 2.2 to add new fields
p18_idx = find_para('项目地址（省/市/区）')
if p18_idx:
    doc.paragraphs[p18_idx].text = (
        '项目地址（省/市/区） | 项目分类（层级路径） | 是否优质项目 | 是否重点区域 | 是否重点领域'
    )

# P19: Update to reflect new field names
p19_idx = find_para('工程类别（原总公司市场口径）')
if p19_idx:
    doc.paragraphs[p19_idx].text = (
        '工程类别 | 重点工程类别 | 业务类型（由项目分类派生） | 是否城市更新 | 城市更新一级分类 | 城市更新二级分类'
    )

# P20: Add new DMP fields
p20_idx = find_para('是否地产类项目 | 是否投资项目')
if p20_idx:
    doc.paragraphs[p20_idx].text = (
        '是否地产类项目 | 是否投资项目 | 是否城市更新 | 是否军融项目 | 关联营销小组'
    )

# P27-31: Update bid/signing section
p27_idx = find_heading('2.4 投标与中标类')
if p27_idx:
    # Find paragraphs in this section
    for i in range(p27_idx+1, p27_idx+10):
        if i < len(doc.paragraphs):
            txt = doc.paragraphs[i].text
            if '招标方式' in txt:
                doc.paragraphs[i].text = (
                    '招标方式 | 评标方法 | 是否联合体投标 | 联合体牵头人性质 | 联合体牵头人 | 资质类型 | 项目承接主体类型'
                )
            elif '交标时间' in txt:
                doc.paragraphs[i].text = (
                    '领取招标文件时间 | 交标时间 | 中标时间 | 中标报量时间 | 预计签约时间 | 中标主体 | 是否签订责任状'
                )
            elif '签约' in txt and '时间' in txt and '中标' not in txt:
                doc.paragraphs[i].text = (
                    '签约时间 | 签约报量时间 | 合同签订年度（由签约时间派生） | 合同性质（执行/备案/一致） | 是否为补充协议'
                )

# P32-36: Update contract amount section
p32_idx = find_heading('2.5 合同金额与效益类')
if p32_idx:
    for i in range(p32_idx+1, p32_idx+8):
        if i < len(doc.paragraphs):
            txt = doc.paragraphs[i].text
            if '签约额' in txt or '中标额' in txt:
                doc.paragraphs[i].text = (
                    '签约额（元） | 自行施工金额（元） | 补充协议金额（元） | 中标额（元，中标报量） | '
                    '项目投资额（万元） | 项目建安合同额（万元） | 预计自行施工造价（万元）'
                )
            elif '效益率' in txt or 'C值' in txt:
                doc.paragraphs[i].text = (
                    '目标效益率（%）（C值） | 管理效益率（%）（B值） | 一次性经营效益率（%）（A值）'
                )

# P37-43: Update payment section
p37_idx = find_heading('2.6 付款与资金类')
if p37_idx:
    for i in range(p37_idx+1, p37_idx+10):
        if i < len(doc.paragraphs) and doc.paragraphs[i].style.name != 'Heading':
            txt = doc.paragraphs[i].text
            if '付款比例' in txt or '进度' in txt:
                doc.paragraphs[i].text = (
                    '月进度付款比例（%） | 节点付款比例（%） | 季度付款比例（%） | 节点时间间隔 | '
                    '进度款付款方式 | 主体完成付款比例（封顶、形象）(%) | 竣工验收支付比例(%) | 结算支付比例（%）'
                )
            elif '非现金' in txt or '垫资' in txt:
                doc.paragraphs[i].text = (
                    '非现金支付比例(%) | 是否垫资 | 垫资比例(%) | 预估垫资金额（万元） | '
                    '累计垫资周期（v2.9新增） | 付款周期（v2.9新增） | 是否付款条件不达标（v2.9新增）'
                )
            elif '预付款' in txt:
                doc.paragraphs[i].text = (
                    '是否有预付款 | 预付款比例（%） | 预付款是否抵扣 | 预付款抵扣方式 | 工程款支付方式'
                )

# P44-49: Update guarantee section
p44_idx = find_heading('2.7 担保类')
if p44_idx:
    for i in range(p44_idx+1, p44_idx+8):
        if i < len(doc.paragraphs) and doc.paragraphs[i].style.name != 'Heading':
            txt = doc.paragraphs[i].text
            if '投标担保' in txt or '投标保证金' in txt:
                doc.paragraphs[i].text = (
                    '投标担保方式 | 投标担保金额（万元） | 汇款银行账号 | 汇款账户名称（v2.9新增） | 汇款银行名称'
                )
            elif '履约担保' in txt:
                doc.paragraphs[i].text = (
                    '履约担保方式 | 履约担保金额（万元） | 履约担保比例（%）'
                )

# P50-53: Update three-certificates section
p50_idx = find_heading('2.8 三证与合规类')
if p50_idx:
    for i in range(p50_idx+1, p50_idx+6):
        if i < len(doc.paragraphs) and doc.paragraphs[i].style.name != 'Heading':
            txt = doc.paragraphs[i].text
            if '规划许可证' in txt or '三证' in txt:
                doc.paragraphs[i].text = (
                    '是否有规划许可证 | 是否有建设用地许可证 | 是否有土地使用证 | '
                    '是否放弃优先受偿权 | 是否限制投标风险（v2.9接入） | 发包方是否无条件禁止承包方停/缓建 | '
                    '承包方因发包方无条件停缓建是否有索赔权利'
                )

# ====================================================================
# 3. UPDATE MODEL 1.4 SECTION - DMP 4-node timeline
# ====================================================================
print('3. Updating Model 1.4 section...')

m14_idx = find_heading('模型1.4：营销统计数据多维度交叉验真模型')
if m14_idx:
    # Find the end of this section (next Heading 2 or Heading 3)
    m14_end = m14_idx + 1
    while m14_end < len(doc.paragraphs):
        if doc.paragraphs[m14_end].style.name.startswith('Heading'):
            break
        m14_end += 1

    # Find paragraphs to update within this section
    for i in range(m14_idx, min(m14_end, m14_idx + 50)):
        p = doc.paragraphs[i]
        txt = p.text

        # Update the data source description
        if '数据来源' in txt and ('DMP' in txt or '中标' in txt or 'OA' in txt):
            p.text = (
                '**数据来源（v2.9更新）**：DMP签约报量（四局）170字段中的时间线字段——'
                '领取招标文件时间、交标时间、中标时间、签约报量时间、签约时间。'
                '这五个字段构成完整的DMP四节点时间线，替代原先依赖OA审批流时间戳的方案。'
                '中标报量导出数据补充中标额（元）、预计签约时间等投标阶段特有字段。'
            )

        # Add new timeline check section before "2. Bid-to-contract cross-validation"
        if 'Bid-to-contract' in txt or '中标额 vs 签约额' in txt or '1a' in txt:
            # Insert description of the 4 new checks
            # We'll insert a paragraph before this one
            pass

    # Find the paragraph right after the data source / methodology intro
    # and insert new timeline content
    insert_after = None
    for i in range(m14_idx, m14_idx + 20):
        if i < len(doc.paragraphs):
            txt = doc.paragraphs[i].text
            if '检测1' in txt or '检测2' in txt or '中标 → 签约报量' in txt or '中标签约金额偏离' in txt:
                insert_after = i - 1
                break
            if '流程校验' in txt or 'Process' in txt or '验真流程' in txt:
                insert_after = i
                break

    # Add v2.9 timeline checks after the methodology intro
    # Find the right insertion point
    if insert_after is None:
        # Find paragraph after heading that describes the data source
        for i in range(m14_idx + 1, m14_idx + 15):
            if '数据来源' in doc.paragraphs[i].text or '检测维度' in doc.paragraphs[i].text:
                insert_after = i + 1
                break

    if insert_after and insert_after < len(doc.paragraphs):
        new_paras = [
            '',
            '**🆕 v2.9新增：DMP四节点时间线交叉校验（替代OA审批流时间戳）**',
            '',
            '签约报量（四局）提供了完整的招投标→签约全链路时间节点，系统新增4项时序倒置检测，',
            '可覆盖原先需OA审批流时间戳才能完成的80%数据验真场景：',
            '',
            '| 检测项 | 字段组合 | 判定逻辑 | 严重等级 | 说明 |',
            '|--------|---------|---------|---------|------|',
            '| 招文领取→交标 | 领取招标文件时间 vs 交标时间 | 领取 > 交标 → 倒置 | 🔴 red | 领取招标文件不应晚于投标截止 |',
            '| 交标→中标 | 交标时间 vs 中标时间 | 交标 > 中标 → 倒置 | 🔴 red | 投标截止不应晚于中标日 |',
            '| 中标→签约报量 | 中标时间 vs 签约报量时间 | 中标 > 报量 → 异常 | 🟡 yellow | 中标结果应在报量之前 |',
            '| 签约报量→签约 | 签约报量时间 vs 签约时间 | 报量 > 签约 → 异常 | 🟡 yellow | 报量不应晚于签约完成 |',
            '',
            '注：OA精确审批时间（bid_review_pass_time、project_approval_time、contract_review_time）',
            '短期仍标记为待采集，DMP四节点可满足当前80%的时序校验需求。',
            '精确的审批流程时间戳可提供更细粒度的流程合规检查。',
            '',
        ]
        # Insert paragraphs after insert_after
        for j, text in enumerate(new_paras):
            para = doc.paragraphs[insert_after + 1 + j]
            # We can't insert paragraphs easily in python-docx with paragraph objects,
            # so we'll use XML manipulation
            pass

# ====================================================================
# 4. UPDATE MODEL 2.1 SECTION - Payment condition cross-validation
# ====================================================================
print('4. Updating Model 2.1 section...')

m21_idx = find_heading('模型2.1：风险分级严禁投标底线检测与审批穿透模型')
if m21_idx:
    # Find end of section
    m21_end = m21_idx + 1
    while m21_end < len(doc.paragraphs):
        if doc.paragraphs[m21_end].style.name.startswith('Heading'):
            break
        m21_end += 1

    # Find a suitable insertion point for the cross-validation description
    for i in range(m21_idx, m21_end):
        txt = doc.paragraphs[i].text
        if '红线规则' in txt or '12项' in txt or '严禁投标' in txt:
            # Add cross-validation after the list of red-line rules
            insert_pt = i
            # Find end of the bullet list
            for j in range(i + 1, min(m21_end, i + 20)):
                if doc.paragraphs[j].text.strip() == '' or '检测逻辑' in doc.paragraphs[j].text:
                    insert_pt = j
                    break

            # We'll add new content by modifying existing empty paragraphs
            # Let's find paragraphs near the end of this section
            for j in range(m21_end - 5, m21_idx, -1):
                if doc.paragraphs[j].text.strip() == '':
                    doc.paragraphs[j].text = (
                        '**🆕 v2.9新增：付款条件不达标标记制度交叉校验**'
                    )
                    if j + 1 < m21_end:
                        doc.paragraphs[j + 1].text = (
                            '签约报量（四局）新增「是否付款条件不达标」字段（col 56），为DMP系统自动标记。'
                            '系统按手册0520/2020版制度规则独立重新计算，与DMP标记交叉比对：'
                        )
                    if j + 2 < m21_end:
                        doc.paragraphs[j + 2].text = (
                            '• DMP标记「否」（达标）但制度校验应不达标 → 🔴 red：'
                            'DMP填列错误或规则未同步，可能导致红线项目漏过。触发条件：月进度付款比例<75%（地产）/70%（非地产）、'
                            '非现金支付比例>12%（地产）/30%（非地产）、存在垫资等。'
                        )
                    if j + 3 < m21_end:
                        doc.paragraphs[j + 3].text = (
                            '• DMP标记「是」（不达标）但制度校验未触发红线 → 🟡 yellow：'
                            'DMP标记依据待核实，可能阈值过严或有其他考量因素。'
                        )
                    if j + 4 < m21_end:
                        doc.paragraphs[j + 4].text = (
                            '• 标记与校验一致 → 不告警，DMP填列正确。'
                        )
                    if j + 5 < m21_end:
                        doc.paragraphs[j + 5].text = (
                            '制度校验覆盖维度：月进度付款比例（区分地产/非地产、手册0520/2020版、橙档/黄绿档）、'
                            '非现金支付比例、是否垫资。校验结果以「付款条件标记校验不一致」问题分类输出，'
                            '违规映射归类至「6_营销统计管理」。'
                        )
                    break

# ====================================================================
# 5. UPDATE CHAPTER 6 - OFFLINE FIELDS
# ====================================================================
print('5. Updating Chapter 6 - offline fields...')

# P790: Chapter title
ch6_idx = find_heading('第六章 模型字段需求总表')
if ch6_idx:
    doc.paragraphs[ch6_idx].text = '第六章 模型字段需求总表（v2.9 精准版）'

# Update the OA section description (P809-811)
oa_idx = find_heading('6.2.6 OA审批流时间戳')
if oa_idx:
    # P811 is the description paragraph after the blank P810
    if oa_idx + 2 < len(doc.paragraphs):
        doc.paragraphs[oa_idx + 2].text = (
            '**状态（v2.9更新）**：pending —— DMP四节点（交标→中标→签约报量→签约）'
            '已可覆盖80%时序校验场景（模型1.4 v2.9新增4项时序倒置检测）。'
            'OA精确审批时间戳（招文评审通过、立项审批、合同评审通过）'
            '仍建议从OA系统导出以提供更细粒度的流程合规检查。'
            'OA共需4个字段，当前优先级：🔴 bid_review_pass_time（招文评审通过时间），'
            '🟡 project_approval_time（立项审批）、contract_review_time（合同评审通过）。'
        )

# Update DMP section (P814-815)
dmp_idx = find_heading('6.2.8 DMP系统增设')
if dmp_idx:
    if dmp_idx + 1 < len(doc.paragraphs):
        doc.paragraphs[dmp_idx + 1].text = (
            '**状态（v2.9更新）**：签约报量（四局）170字段已确认含23个新增关键字段，'
            '包括三证（规划/建设/土地使用许可证）、DMP四节点时间线（交标→中标→签约报量→签约）、'
            '是否付款条件不达标、是否限制投标风险、是否为补充协议、合同性质（执行/备案/一致）、'
            '累计垫资周期、付款周期、城市更新一级/二级分类、联合体牵头人性质、是否重点领域、'
            '关联营销小组、造价咨询单位、其他资金来源说明、汇款账户名称等。'
            '仅「是否需协助业主融资」1个字段仍待DMP补充。'
        )

# Update summary paragraph P818
for i in range(815, 825):
    if i < len(doc.paragraphs):
        txt = doc.paragraphs[i].text
        if '当前最急缺' in txt:
            doc.paragraphs[i].text = (
                '**当前最急缺（v2.9）**：SAP 4个日期字段（模型2.3核心）。'
                '客户台账DMP中标报量已覆盖全部维度三需求；OA审批流DMP四节点已覆盖80%场景，精确时间仍建议获取。'
                '共 **4个🔴高优字段**（较v2.8减少5个，客户台账3字段+OA下调）。'
            )
            break

# ====================================================================
# 6. UPDATE CHAPTER 7 - RECOMMENDATIONS
# ====================================================================
print('6. Updating Chapter 7...')

ch7_idx = find_heading('第七章 补充建议与完善意见')
if ch7_idx:
    # Find "建议增加的数据源" and update
    for i in range(ch7_idx, ch7_idx + 30):
        if i < len(doc.paragraphs):
            txt = doc.paragraphs[i].text
            if 'OA审批' in txt and ('时间' in txt or '流程' in txt):
                doc.paragraphs[i].text = (
                    '**OA审批流时间戳（v2.9状态：DMP四节点已可覆盖80%场景）**：'
                    '交标时间、中标时间、签约报量时间、签约时间已在模型1.4中实现全套时序校验。'
                    'OA精确审批时间仍建议获取以便更完整的流程合规审计。'
                )

# Update Chapter 7 recommendations
# P840: Update 住建部 data note
p840_idx = find_para('住建部四库一平台数据')
if p840_idx:
    doc.paragraphs[p840_idx].text = (
        '2. **住建部四库一平台数据**：获取项目三证信息、企业资质信息 → '
        '模型2.1三证合规自动核验。**v2.9更新**：签约报量（四局）已含三证字段'
        '（是否有规划/建设/土地使用许可证），不再依赖附表1手工填报三证数据。'
        '但四库一平台可作为交叉验证的独立数据源。'
    )

# P841: Update OA note
p841_idx = find_para('OA审批流时间戳数据')
if p841_idx:
    doc.paragraphs[p841_idx].text = (
        '3. **OA审批流时间戳数据**：投标立项审批时间、招文评审审批时间、合同评审审批时间 → '
        '模型1.4未批先投检测、模型2.1审批穿透检测。**v2.9更新**：DMP四节点时间线'
        '（交标→中标→签约报量→签约）已在模型1.4中实现4项时序倒置检测，覆盖80%场景。'
        'OA精确审批时间仍建议获取用于更完整的审批穿透审计。'
    )

# Add a v2.9 changelog at the end
print('7. Adding revision changelog section...')

# Find the last appendix heading
last_appendix_idx = None
for i, p in enumerate(doc.paragraphs):
    if '补充说明' in p.text and p.style.name.startswith('Heading'):
        last_appendix_idx = i
        break

if last_appendix_idx:
    # Add changelog paragraphs at the very end
    # Find the last non-empty paragraph
    end_idx = len(doc.paragraphs) - 1
    while end_idx > 0 and doc.paragraphs[end_idx].text.strip() == '':
        end_idx -= 1

    # We need to add content after the last paragraph
    # In python-docx, we add paragraphs to the document body
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Add a new section with heading
    body = doc.element.body

    def add_heading_para(text, level=1):
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), f'Heading{level}')
        pPr.append(pStyle)
        p.append(pPr)
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        p.append(r)
        body.append(p)

    def add_body_para(text):
        p = OxmlElement('w:p')
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        p.append(r)
        body.append(p)

    # Add changelog section
    add_body_para('')
    add_body_para('')
    add_heading_para('附：v2.9 版本修订明细', level=1)
    add_body_para('')
    add_body_para('以下为v2.8→v2.9的全部修订内容：')
    add_body_para('')

    changes = [
        ('模型1.4 — DMP四节点时间线校验',
         '新增4项时序倒置检测：招文领取→交标（red）、交标→中标（red）、'
         '中标→签约报量（yellow）、签约报量→签约（yellow）。'
         '使用DMP签约报量（四局）字段：领取招标文件时间、交标时间、中标时间、签约报量时间、签约时间。'
         '替代原先依赖OA审批流时间戳的方案。'),
        ('模型2.1 — 付款条件不达标制度交叉校验',
         '新增「是否付款条件不达标」字段的独立制度验证。按手册0520/2020版规则重新计算，'
         '与DMP系统标记交叉比对。DMP误标为达标但实际应不达标 → red；'
         'DMP误标为不达标但实际达标 → yellow。校验维度：月进度付款比例、非现金支付比例、垫资。'),
        ('dmp_loader — 数值转换补充',
         '新增累计垫资周期、付款周期两个字段的safe_float数值转换。'),
        ('代码审查修复',
         'model_2_2_profit.py：修复zombie_days未定义（NameError），从experience_warnings配置中提取。'
         'report_generator.py：版本号v2.8→v2.9；摘要key「僵尸合同」→「停工退场停缓建预警」+「在施状态存疑」。'
         'supervision_compliance.py：新增4项时间线倒置和付款标记校验不一致的违规映射。'),
        ('offline_fields — 40→63字段',
         '新增23个DMP字段：三证（3）、时间线（5，含领取招标文件时间）、付款条件标记、'
         '限制投标风险、补充协议标记、合同性质、累计垫资周期、付款周期、汇款账户名称、'
         '城市更新分类（2）、联合体牵头人性质、是否重点领域、关联营销小组、造价咨询单位、其他资金来源说明。'
         'OA审批流时间戳标注为"DMP四节点可部分替代"。'),
        ('数据源适配 — 签约报量（四局）',
         '确认签约报量（四局）170字段为DMP新版导出格式，使用新版命名（优质客户/重点区域/优质项目/合同性质）。'
         '中标报量保留旧版命名（高端客户/市场/项目/项目模式类），作为投标阶段补充数据源。'),
    ]

    for title, desc in changes:
        add_body_para(f'**{title}**')
        add_body_para(desc)
        add_body_para('')

# ====================================================================
# SAVE
# ====================================================================
print(f'Saving to: {DST}')
doc.save(DST)
print('Done!')
